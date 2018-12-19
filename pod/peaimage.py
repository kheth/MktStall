from operator import attrgetter
import PIL
import cv2
import docx
import pytesseract
from PIL import Image
from lxml import etree
from shapely.geometry import Polygon
from skimage import color
from skimage import feature
from skimage import io  # image I/O
from skimage import measure  # import measure
from skimage.measure import regionprops  # measure step
from skimage.morphology import disk, dilation
from skimage.transform import rescale  # hsv converting
import imagelabel
import math
import os
from parameters import Parameters
from peautils import PeaUtils
from rulermeasure import RulerMeasure
from configmaker import ConfigMaker
import sys
from matplotlib import pyplot as plt
import numpy as np
import csv
from scipy import signal
from scipy.signal import savgol_filter
from PIL import ImageDraw
from ast import literal_eval as make_tuple
from shapely.geometry import Point
from shapely.geometry.polygon import Polygon

#############################
# File extensions
# A = rotated transformation
# B = greyscale transformation
# C = adaptive threshold transformation
# D = crop transformation
# E = denoise the image
#############################
# F = global_thresholding
# G = otsus_thresholding
# H = gaussian_filtering_otsus_thresholding
# im.size # (width,height) tuple
#############################


class PeaImage:

    # Type of Qualitative Image Assessment
    TypeGoodImage = "TYPE_GOOD_IMAGE"
    TypeBadImage = "TYPE_BAD_IMAGE"
    TypeSlightProblemImage = "TYPES_LIGHT_PROBLEM_IMAGE"
    TypeDuplicateImage = "TYPE_DUPLICATE_IMAGE"

    def __init__(self, name, path, config):
        self.DEBUG = True
        self.VERBOSE = True
        self.print_contour_points_to_std_out = False
        self.SHOW_IMAGES = False
        self.name = name
        self.path = path
        self.image_label = ""
        self.image_ruler = RulerMeasure("nothing", config)
        self.joined_final_l = "" # this is your label separated by underscore
        self.jicAccNo = ""
        self.structureCO = ""
        self.uniqueID = ""
        self.group = ""
        self.subGroup = ""
        self.replicate = ""
        self.format = ""
        self.format_conversion = ""
        self.resolution = ""
        self.dimension = ""
        self.orientation = ""
        self.rotated = ""
        self.createDate = ""
        self.colourSpace = ""
        self.fileSize = ""
        # Qualitative Image Assessment
        self.imageAssessment = ""
        # List the objects that comprise the image
        self.objects = []
        self.number_of_teeth = ""
        self.a_magic_number_for_inequality = 0.7

        # IMAGE RESIZING PARAMETERS
        # 50% more pixels (less pixels!)
        self.resize_value = 0.7
        # Magic number for leaf tooth area
        self.magic_number_from_paper = 15
        # std deviation multiplier
        self.std_multiplier = 1.5
        self.magic_height_number = 4.0
        self.std_dev_length_multiplier = 0.5

        # multiplyer for median
        self.hue_median_multiplier = 0.925

        self.hsv_value = ""
        self.hue_value = ""
        self.shape = ""
        self.mean_hsv = ""
        self.mean_hue_value = ""
        self.median_hue_value = ""
        self.scale_ROI = ""
        self.edge_detection = ""
        self.refined_ROI = ""
        self.newRegionsSorted = ""
        self.leaf_tmp = ""
        self.contour_tmp = ""
        self.intersects = ""
        self.ref_coord = ""
        self.leaf_tmp = ""
        self.contours = ""
        self.imageAcceptable = False
        self.accepted_labels = []
        self.cropped_region_image_object = ""
        self.csv_string = ""
        self.distances = []
        self.new_position_index_list = []
        self.mean_dist_from_centroid = []
        self.onestd_multiplierabove = []
        self.onestd_multiplierbelow = []
        self.one_std_above = []
        self.one_std_below = []
        self.two_std_above = []
        self.two_std_below = []
        self.config_file = ConfigMaker(self.path)

    def get_bordered_leaf_full_image(self):

        src = cv2.imread(self.path, cv2.CV_LOAD_IMAGE_COLOR)
        bgr = cv2.split(src)
        # Note OpenCV uses BGR color order
        cv2.imwrite(self.path + ".blue.jpg", bgr[0]);  # blue channel
        cv2.imwrite(self.path + ".green.jpg", bgr[1]);  # green channel
        cv2.imwrite(self.path + ".red.jpg", bgr[2]);  # red channel

        img = cv2.imread(self.path + ".blue.jpg", 0)
        ret, thresh1 = cv2.threshold(img, 180, 255, cv2.THRESH_BINARY)

        save_path = self.path + ".new_thresh.jpg"
        cv2.imwrite(save_path, thresh1)

        #################################
        img = Image.open(save_path)
        im = Image.new("RGB", img.size)
        pixels = img.load()  # create the pixel map
        for i in range(img.size[0]):  # for every pixel:
            for j in range(img.size[1]):
                #print "pixel i", i, "j", j, "val", pixels[i,j]
                if pixels[i, j] <= 10:
                    im.putpixel((i, j), (255, 240, 11))
                else:
                    im.putpixel((i,j), (20, 100, 255))
                    #pixels[i, j] = (255, 240, 11)

        path = self.path + ".20000001.jpg"
        im.save(path)
        #################################

        new_im = io.imread(path)
        im_resize = rescale(new_im, 2.0, mode='constant')
        path = self.path + ".20000002.jpg"
        io.imsave(path, im_resize)


        self.config_file.original_image_output_config_file(self.path)
        sizing_image_only = Image.open(path)
        dimensions = sizing_image_only.size
        width = dimensions[0] * self.resize_value
        height = dimensions[1] * self.resize_value

        reference_image = io.imread(path)
        # Pre-processing of images by rescaling and converting to HSV
        # Step 2a: Pre-processing of images by rescaling and converting to HSV
        pea_resize = rescale(reference_image, self.resize_value, mode='constant')  # 50% more pixels
        pea_hsv = color.rgb2hsv(pea_resize)  # convert to hsv
        # every thing in hue dimension, everything in sat, only 2nd index in brightness value
        pea_value = pea_hsv[:, :, 2]
        # every thing in hue dimension, everything in sat, only 0th index in brightness value
        pea_hue = pea_hsv[:, :, 0]

        # Step 2b: print and check pre-processing - no transformations in this block!!!
        # find geometry of shape, x by y in pixels, and how many hsv values each pixel has

        # H + S + V /3 for each pixel and then average for all pixels!
        # find mean aggregate HSV of resized and valued extracted pea shapes

        # find median aggregate HSV of resized and valued extracted
        #print "Forming Pea_Resize and Pea_Hue."
        # fig, (ax1, ax2) = plt.subplots(ncols=2, nrows=1, figsize=(20, 15))
        #  show 2 pictures with width/height in inches
        # ax1.imshow(Pea_Resize)  # resize picture
        # ax2.imshow(Pea_Hue, cmap='flag')  # hue as a gray colourmap
        # ax2.imshow(Pea_Hue)  # hue as a gray colourmap
        # os.remove(imageNameTemp)  # delete

        # Step 2c: Thresholding
        # Value thresholding using median intensity value what becomes white/black
        # if pixel colour is above median value is one/true then white
        # if pixel colour is below median value is zero/false then black
        # Scale_ROI is true or false of Pea_Hue matrix
        # 0.925 Hard-coded, should use Python Stats to calculate
        # Scale_ROI = Pea_Hue > np.median(Pea_Hue) * 0.925
        scale_roi = pea_hue > np.median(pea_hue) * self.hue_median_multiplier

        # Step 2d: Detect Leaf Edges
        # canny edge detection to find the outline of the leaf
        # input greyscale image, find the edge and do not smooth with sigma
        # dilation takes the pixels of the edge and sets its neighbour to be brighter,
        # disk takes 1 pixel in each direction, changes thicknes
        # edges_detection = feature.canny(color.rgb2gray(Pea_Hue), sigma= 0) gives a less clean image
        # dilated_Edge = dilation(edges_detection, disk(1))
        # edges_detection = feature.canny(color.rgb2gray(pea_hue), sigma=3)
        #feature.canny(color.rgb2gray(pea_hue), sigma=3)  # type: object
        # dilated_Edge = dilation(edges_detection, disk(3))

        # Step 2d: print and check thresholding and dilation
        print "Forming logical Scale ROI and dilated Image."
        # fig, (ax1, ax2) = plt.subplots(ncols=2, nrows=1, figsize=(20, 15))
        # ax1.imshow(np.logical_not(Scale_ROI), cmap='gray')
        # ax1.imshow(np.logical_not(Scale_ROI))
        # ax2.imshow(dilated_Edge, cmap='gray')

        # Step 2e: Find XOR
        # create inverse of Scale ROI
        refined_roi = np.logical_not(scale_roi)
        # draw white if the two pictures above differ, black if they are the same
        print "Forming logical xor of refined ROI and dilated edge."
        # plt.figure(figsize=(15, 15))
        # plt.imshow(np.logical_xor(Refined_ROI, dilated_Edge), cmap='gray')

        # Step 2f: Find objects
        # Close objects to fill holes in detected objects
        # Find contours at a constant value of 0.9
        leaf_mask = refined_roi

        # Generate the cubic framework
        # How many neighbours does it have with only connected by 2
        label_img2 = measure.label(leaf_mask, connectivity=2)  # 2D matrix of connected pixels
        # of this, begin making measurements
        # co-ordinates of connected pixels and value
        regions = regionprops(label_img2, intensity_image=pea_hsv[:, :, 2])
        # Set empty array and variables for storing measurements SECOND PEA PROBLEM STARTS
        blank_img = np.zeros((pea_hue.shape[0], pea_hue.shape[1]), dtype=np.uint8)
        # largest_area = 0
        leaf_tmp = blank_img.copy()  # also another array of zeros
        # contour_tmp = blank_img.copy()  # also another array of zeros

        ####################
        # figure_grey_clean = self.get_greyscale_copy("figure_grey_clean", self.path)
        # figure_input = io.imread(self.path, as_grey=True, flatten=True)  # type: object
        # figure_grey_clean_resize = rescale(figure_input, self.resize_value, mode='constant')  # 50% more pixels
        ####################

        # fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(20, 20))
        # ax.imshow(figure_grey_clean_resize, cmap='gray')  # no blue line
        # creates a new sorted region by area
        new_regions_sorted = sorted(regions, key=attrgetter('area'), reverse=True)

        # Step 2g: Choose the leaf and measure it
        # iterate thorough list of ordered by size regions and take the first region found inside the bounding box
        # will be the largest region and will be the leaf
        # draw blue line around leaf
        for region in new_regions_sorted:

            intersects = PeaUtils.point_within_hardcode_bounding_box(width,
                                                                     height,
                                                                     region.bbox[0],
                                                                     region.bbox[1],
                                                                     region.bbox[2],
                                                                     region.bbox[3])
            if intersects != "true":
                continue

            ##################################################
            # This block populates and displays Leaf_tmp
            # coordinates based on labelled outlines, draws blue lines
            ref_coord = region.coords.astype(int)
            leaf_tmp[ref_coord[:, 0], ref_coord[:, 1]] = 1

            io.imsave(self.path + ".2.jpg", pea_resize)
            pea_resize_outline = Image.open(self.path + ".2.jpg")
            #fully_connected='low', positive_orientation='low'
            contours = measure.find_contours(leaf_tmp, 0.8, fully_connected='high')
            keep_printing = True
            counter = 0
            for n, contour in enumerate(contours):
                #print "n: ", n
                #print "length: ", len(contour)
                if len(contour) > 200:  # 201 is the minimum number of pixels in a contour for a contour to be drawn.
                    #print "drawing"
                    for m in enumerate(contour):
                        value = (0, 0, 0)
                        xy = (int(round(m[1][1])), int(round(m[1][0])))
                        pea_resize_outline.putpixel(xy, value)
                        #if keep_printing == True:
                            #print "CONTOUR", xy
                            #counter += 1
                            #if counter == 10:
                            #    keep_printing = False

            io.imsave(self.path + ".2.jpg", pea_resize_outline)
        return self.path + ".2.jpg"

    def analyse_image(self):

        #plt.close('all')
        path = self.get_bordered_leaf_full_image()
        self.config_file.original_image_output_config_file(path)
        sizing_image_only = Image.open(path)
        dimensions = sizing_image_only.size
        width = dimensions[0] * self.resize_value
        height = dimensions[1] * self.resize_value

        reference_image = io.imread(path)
        # Pre-processing of images by rescaling and converting to HSV
        # Step 2a: Pre-processing of images by rescaling and converting to HSV
        pea_resize = rescale(reference_image, self.resize_value, mode='constant')  # 50% more pixels
        pea_hsv = color.rgb2hsv(pea_resize)  # convert to hsv
        # every thing in hue dimension, everything in sat, only 2nd index in brightness value
        pea_value = pea_hsv[:, :, 2]
        # every thing in hue dimension, everything in sat, only 0th index in brightness value
        pea_hue = pea_hsv[:, :, 0]


        # Step 2b: print and check pre-processing - no transformations in this block!!!
        # find geometry of shape, x by y in pixels, and how many hsv values each pixel has

        # H + S + V /3 for each pixel and then average for all pixels!
        # find mean aggregate HSV of resized and valued extracted pea shapes

        # find median aggregate HSV of resized and valued extracted
        print "Forming Pea_Resize and Pea_Hue."
        # fig, (ax1, ax2) = plt.subplots(ncols=2, nrows=1, figsize=(20, 15))
        #  show 2 pictures with width/height in inches
        # ax1.imshow(Pea_Resize)  # resize picture
        # ax2.imshow(Pea_Hue, cmap='flag')  # hue as a gray colourmap
        # ax2.imshow(Pea_Hue)  # hue as a gray colourmap
        # os.remove(imageNameTemp)  # delete

        # Step 2c: Thresholding
        # Value thresholding using median intensity value what becomes white/black
        # if pixel colour is above median value is one/true then white
        # if pixel colour is below median value is zero/false then black
        # Scale_ROI is true or false of Pea_Hue matrix
        # 0.925 Hard-coded, should use Python Stats to calculate
        # Scale_ROI = Pea_Hue > np.median(Pea_Hue) * 0.925
        scale_roi = pea_hue > np.median(pea_hue) * self.hue_median_multiplier

        # Step 2d: Detect Leaf Edges
        # canny edge detection to find the outline of the leaf
        # input greyscale image, find the edge and do not smooth with sigma
        # dilation takes the pixels of the edge and sets its neighbour to be brighter,
        # disk takes 1 pixel in each direction, changes thicknes
        # edges_detection = feature.canny(color.rgb2gray(Pea_Hue), sigma= 0) gives a less clean image
        # dilated_Edge = dilation(edges_detection, disk(1))
        #edges_detection = feature.canny(color.rgb2gray(pea_hue), sigma=3)
        feature.canny(color.rgb2gray(pea_hue), sigma=3)  # type: object
        #dilated_Edge = dilation(edges_detection, disk(3))


        # Step 2d: print and check thresholding and dilation
        print "Forming logical Scale ROI and dilated Image."
        # fig, (ax1, ax2) = plt.subplots(ncols=2, nrows=1, figsize=(20, 15))
        # ax1.imshow(np.logical_not(Scale_ROI), cmap='gray')
        # ax1.imshow(np.logical_not(Scale_ROI))
        # ax2.imshow(dilated_Edge, cmap='gray')



        # Step 2e: Find XOR
        # create inverse of Scale ROI
        refined_roi = np.logical_not(scale_roi)
        # draw white if the two pictures above differ, black if they are the same
        print "Forming logical xor of refined ROI and dilated edge."
        # plt.figure(figsize=(15, 15))
        # plt.imshow(np.logical_xor(Refined_ROI, dilated_Edge), cmap='gray')

        # Step 2f: Find objects
        # Close objects to fill holes in detected objects
        # Find contours at a constant value of 0.9
        leaf_mask = refined_roi
        # Generate the cubic framework
        # How many neighbours does it have with only connected by 2
        label_img2 = measure.label(leaf_mask, connectivity=2)  # 2D matrix of connected pixels
        # of this, begin making measurements
        # co-ordinates of connected pixels and value
        regions = regionprops(label_img2, intensity_image=pea_hsv[:, :, 2])
        # Set empty array and variables for storing measurements SECOND PEA PROBLEM STARTS
        blank_img = np.zeros((pea_hue.shape[0], pea_hue.shape[1]), dtype=np.uint8)
        # largest_area = 0
        leaf_tmp = blank_img.copy()  # also another array of zeros
        # contour_tmp = blank_img.copy()  # also another array of zeros

        ####################
        # figure_grey_clean = self.get_greyscale_copy("figure_grey_clean", self.path)
        #figure_input = io.imread(self.path, as_grey=True, flatten=True)  # type: object
        #figure_grey_clean_resize = rescale(figure_input, self.resize_value, mode='constant')  # 50% more pixels
        ####################

        #fig, ax = plt.subplots(ncols=1, nrows=1, figsize=(20, 20))
        #ax.imshow(figure_grey_clean_resize, cmap='gray')  # no blue line
        # creates a new sorted region by area
        new_regions_sorted = sorted(regions, key=attrgetter('area'), reverse=True)

        # Step 2g: Choose the leaf and measure it
        # iterate thorough list of ordered by size regions and take the first region found inside the bounding box
        # will be the largest region and will be the leaf
        # draw blue line around leaf
        for region in new_regions_sorted:

            intersects = PeaUtils.point_within_hardcode_bounding_box(width,
                                                                    height,
                                                                    region.bbox[0],
                                                                    region.bbox[1],
                                                                    region.bbox[2],
                                                                    region.bbox[3])
            if intersects != "true":
                continue

            ##################################################
            # This block populates and displays Leaf_tmp
            # coordinates based on labelled outlines, draws blue lines
            ref_coord = region.coords.astype(int)
            leaf_tmp[ref_coord[:, 0], ref_coord[:, 1]] = 1

            #io.imsave(self.path + ".2.jpg", pea_resize)
            #pea_resize_outline = Image.open(self.path + ".2.jpg")
            #contours = measure.find_contours(leaf_tmp, 0.9)
            #for n, contour in enumerate(contours):
            #     for m in enumerate(contour):
            #         value = (12,43, 18)
            #         xy = (int(round(m[1][1])),int(round(m[1][0])))
            #         pea_resize_outline.putpixel(xy, value)
            #         print "CONTOUR", xy
            #    ax.plot(contour[:, 1], contour[:, 0], '-b', linewidth=1)
            # row,col
            #io.imsave(self.path + ".2.jpg", pea_resize_outline)

            y0, x0 = region.centroid
            orientation = region.orientation
            x1 = x0 + math.cos(orientation) * 0.5 * region.major_axis_length  # col
            y1 = y0 - math.sin(orientation) * 0.5 * region.major_axis_length  # row
            x2 = x0 - math.cos(orientation) * 0.5 * region.major_axis_length  # col
            y2 = y0 + math.sin(orientation) * 0.5 * region.major_axis_length  # row
            # ax.plot((x1, x2), (y1, y2), '-r', linewidth=1)
            #ax.plot(x0, y0, '.g', markersize=5)

            coords_list = region.coords
            y_list = []
            x_list = []
            for t in coords_list:
                y_list.append(t[0])
                x_list.append(t[1])
            min_x = min(x_list) - 80
            min_y = min(y_list) - 80
            max_x = max(x_list) + 80
            max_y = max(y_list) + 80
            #ax.plot((min_x, max_x), (min_y, min_y), '-r', linewidth=1)
            #ax.plot((max_x, max_x), (min_y, max_y), '-g', linewidth=1)
            #ax.plot((max_x, min_x), (max_y, max_y), '-b', linewidth=1)
            #ax.plot((min_x, min_x), (max_y, min_y), '-c', linewidth=1)

            crop_img = pea_resize[min_y:max_y, min_x:max_x]  # Crop from x, y, w, h -> 100, 200, 300, 400
            # NOTE: its img[y: y + h, x: x + w] and *not* img[x: x + w, y: y + h]
            io.imsave(path + ".cr.jpg", crop_img)

            cropped_region_image_object = PeaImage("cropped_region", path + ".cr.jpg", self.config_file)
            # add success for this process to the config
            self.config_file.crop1_image_output_config_file(path + ".cr.jpg")
            # is true
            break
            # end of choosing for loop
        # assign the first cropped region to self.
        self.cropped_region_image_object = cropped_region_image_object
        # start to measure the initial cropped image.
        self.measure_image(cropped_region_image_object)

    @staticmethod
    def orient_image(cropped_region_image_object, config_file):
        path = cropped_region_image_object.path
        image = cv2.imread(path)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.bilateralFilter(gray, 11, 17, 17)
        edged = cv2.Canny(gray, 30, 200)
        # find contours in the edged image, keep only the largest
        # ones, and initialize our screen contour. Fit an ellipse to the contour
        # then rotate the image.
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (15, 15))
        dilated = cv2.dilate(edged.copy(), kernel)
        (contours_found, hierarchy) = cv2.findContours(dilated.copy(), cv2.RETR_TREE, cv2.CHAIN_APPROX_NONE)
        contours_found = sorted(contours_found, key=cv2.contourArea, reverse=True)[:10]
        contour = contours_found[0]
        (x, y), (MA, ma), angle = cv2.fitEllipse(contour)
        # ASSUMPTION: The top of the leaf is in the range of -90 to 90 degrees.
        # If its not, then fitEllipse has miss identified the top.
        if 90 < angle < 270:
            if 90 < angle <= 180:
                angle = -(180.0 - angle)
            if 180 < angle <= 270:
                angle = angle - 180
        rows, cols, channels = image.shape
        m = cv2.getRotationMatrix2D((x, y), angle, 1)
        rotated_img = cv2.warpAffine(image, m, (cols, rows))
        rotated_img[np.where((rotated_img <= [1, 1, 1]).all(axis=2))] = [251, 106, 29]
        cv2.imwrite(path + ".ro.jpg", rotated_img)
        rotated_image_object = PeaImage("rotated_image", path + ".ro.jpg", config_file)
        return rotated_image_object

    def get_canny_edges(self, image_object):
        path = image_object.path
        image = cv2.imread(path)
        #cv2.imwrite(path + ".a1.IMAGE.jpg", image)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        #cv2.imwrite(path + ".a1.GRAY.jpg", gray)

        gray = cv2.bilateralFilter(gray, 17, 17, 17)
        edged = cv2.Canny(gray, 0, 200)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated1 = cv2.dilate(edged.copy(), kernel)
        #cv2.imwrite(path + ".a1.DILATED_1.jpg", dilated1)

        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.bilateralFilter(gray, 20, 17, 17)
        edged = cv2.Canny(gray, 0, 200)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated2 = cv2.dilate(edged.copy(), kernel)
        #cv2.imwrite(path + ".a1.DILATED_2.jpg", dilated2)

        dilated_or_1 = cv2.bitwise_or(dilated1, dilated2)
        #cv2.imwrite(path + ".a1.DILATED_1and2.jpg", dilated_or_1)

        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.bilateralFilter(gray, 20, 30, 50)
        edged = cv2.Canny(gray, 0, 500)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated3 = cv2.dilate(edged.copy(), kernel)
        #cv2.imwrite(path + ".a1.DILATED_3.jpg", dilated3)

        dilated_or_2 = cv2.bitwise_or(dilated_or_1, dilated3)
        #cv2.imwrite(path + ".a1.DILATED_1and2and3.jpg", dilated_or_2)

        edged = cv2.Canny(image, 0, 500)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated4 = cv2.dilate(edged.copy(), kernel)
        #cv2.imwrite(path + ".a1.DILATED_4.jpg", dilated4)

        dilated_or_3 = cv2.bitwise_or(dilated_or_2, dilated4)
        #cv2.imwrite(path + ".a1.DILATED_1and2and3and4.jpg", dilated_or_3)

        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.bilateralFilter(gray, 20, 30, 50)
        edged = cv2.Canny(gray, 10, 250)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated5 = cv2.dilate(edged.copy(), kernel)
        #cv2.imwrite(path + ".a1.DILATED_5.jpg", dilated5)

        dilated_or_4 = cv2.bitwise_or(dilated_or_3, dilated5)
        cv2.imwrite(path + ".a1.DILATED_1and2and3.jpg", dilated_or_4)

        cnts = PeaImage("cnts", path + ".a1.DILATED_1and2and3.jpg", self.config_file)
        return cnts

    def get_leaf_mask(self, path, ID):
        pea_resize = io.imread(path)
        # Pre-processing of images by rescaling and converting to HSV
        # Step 2a: Pre-processing of images by rescaling and converting to HSV
        # pea_resize = rescale(reference_image, 1, mode='constant')  # not necessary...
        pea_hsv = color.rgb2hsv(pea_resize)  # convert to hsv
        pea_hue = pea_hsv[:, :, 0]
        scale_roi = pea_hue > np.median(pea_hue) * self.hue_median_multiplier
        refined_roi = np.logical_not(scale_roi)
        leaf_mask = refined_roi
        label_img2 = measure.label(leaf_mask, connectivity=2)
        regions = regionprops(label_img2, intensity_image=pea_hsv[:, :, 2])
        blank_img = np.zeros((pea_hue.shape[0], pea_hue.shape[1]), dtype=np.uint8)
        leaf_tmp = blank_img.copy()
        new_regions_sorted = sorted(regions, key=attrgetter('area'), reverse=True)

        # for region in new_regions_sorted:
        region = new_regions_sorted[0]
        ref_coord = region.coords.astype(int)
        leaf_tmp[ref_coord[:, 0], ref_coord[:, 1]] = 1
        io.imsave(path + ".1_mask_pre.jpg", leaf_tmp)
        testImg = Image.open(path + ".1_mask_pre.jpg")
        contours = measure.find_contours(leaf_tmp, 0.9)
        for n, contour in enumerate(contours):
            for m in enumerate(contour):
                value = (0, 0, 0)
                xy = (int(round(m[1][1])), int(round(m[1][0])))
                testImg.putpixel(xy, 255)
                #print "CONTOUR", xy
        io.imsave(path + ".2_mask_post." + ID + ".jpg", testImg)
        mask = PeaImage("mask_post", path + ".2_mask_post." + ID + ".jpg", self.config_file)
        return mask

    def crop_rotated_image(self, rotated_image_object):
        path = rotated_image_object.path
        mask = self.get_leaf_mask(path, "cr_ro_cr2")
        leafe_outline = io.imread(mask.path)

        # find contours
        (contours_found, hierarchy) = cv2.findContours(leafe_outline.copy(), cv2.RETR_TREE, cv2.CHAIN_APPROX_NONE)
        contours_found = sorted(contours_found, key=cv2.contourArea, reverse=True)[:10]

        image = cv2.imread(path)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        cv2.drawContours(gray, contours_found[0], -1, (0, 255, 0), 2)
        cv2.imwrite(path + ".3_mask_post_contour.jpg", gray)

        contour = contours_found[0]

        # find extreme left, right, top and bottom points.
        leftmost = tuple(contour[contour[:, :, 0].argmin()][0])
        rightmost = tuple(contour[contour[:, :, 0].argmax()][0])
        topmost = tuple(contour[contour[:, :, 1].argmin()][0])
        bottommost = tuple(contour[contour[:, :, 1].argmax()][0])
        # crop the rotated image
        min_y = topmost[1] - 30
        max_y = bottommost[1] + 30
        min_x = leftmost[0] - 30
        max_x = rightmost[0] + 30
        image = image[min_y:max_y, min_x:max_x]
        cv2.imwrite(path + ".c2.jpg", image)
        cropped_rotated_image_object = PeaImage("cropped2_rotated_image", path + ".c2.jpg", self.config_file)
        return cropped_rotated_image_object

    def measure_image(self, image):
        ##################
        rotated_img_object = self.orient_image(image, self.config_file)
        self.config_file.rotate_image_output_config_file(self.path + ".cr.jpg.ro.jpg")
        cropped_rotated_img_object = self.crop_rotated_image(rotated_img_object)
        self.config_file.crop2_image_output_config_file(self.path + ".cr.jpg.ro.jpg.c2.jpg")
        path = cropped_rotated_img_object.path
        image = cv2.imread(path)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        gray = cv2.bilateralFilter(gray, 11, 17, 17)
        edged = cv2.Canny(gray, 20, 210)

        # show img with canny edge detection
        #fig_edges = plt.figure(figsize=(10, 15))
        #im1 = fig_edges.add_subplot(131)  # left
        #im1.imshow(edged)
        #im1.set_title("edged")

        ###
        # complete contours
        # kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (1, 2))
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated = cv2.dilate(edged, kernel)
        # show img with dilated edges on kanny edge detection
        #im2 = fig_edges.add_subplot(132)  # middle
        #im2.imshow(dilated)
        #im2.set_title("dilated")
        ###########################

        #rotated_img_object = self.orient_image(image, self.config_file)
        #self.config_file.rotate_image_output_config_file(self.path + ".cr.jpg.ro.jpg")
        #cropped_rotated_img_object = self.crop_rotated_image(rotated_img_object)
        #self.config_file.crop2_image_output_config_file(self.path + ".cr.jpg.ro.jpg.c2.jpg")

        #path = cropped_rotated_img_object.path
        #image = cv2.imread(path)
        #gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)


        # show img with canny edge detection
        #fig_edges = plt.figure(figsize=(10, 15))
        #im1 = fig_edges.add_subplot(131)  # left
        #im1.imshow(edged)
        #im1.set_title("edged")

        ###
        # complete contours
        # kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (1, 2))
        #kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        #dilated = cv2.dilate(edged, kernel)

        #eroded = cv2.erode(dilated.copy(), kernel)
        #cv2.imwrite(path + ".A1_TESTING_CONTOUR_AGRAY.jpg", gray)
        #cv2.imwrite(path + ".A1_TESTING_CONTOUR_CANNY.jpg", edged)
        #cv2.imwrite(path + ".A1_TESTING_CONTOUR_DILATED.jpg", dilated)
        #cv2.imwrite(path + ".A1_TESTING_CONTOUR_ERODED.jpg", eroded)

        # show img with dilated edges on canny edge detection
        #im2 = fig_edges.add_subplot(132)  # middle
        #im2.imshow(dilated)
        #im2.set_title("dilated")
        new_contours, hierarchy = cv2.findContours(dilated.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
        contours = sorted(new_contours, key=cv2.contourArea, reverse=True)[:10]
        contour = contours[0]

        ###################################################
        if self.print_contour_points_to_std_out:
            single_contour = contours[0]
            Pointbegin = single_contour[0]
            length = len(single_contour)
            Pointend = single_contour[length - 1]

            print "cSTART: ", Pointbegin
            print "cEND: ", Pointend
            print "START - ################"
            prev = contour[0]
            for x in contour:
                #if x[0] < prev[0]-1 or x[0] > prev[0]+1:
                print "NOT NORMAL\t", x[0][0], "\t", x[0][1]
                prev = x
            print "END - ################"
        #####################################################

        # make skinny contours for teeth finding later in the code.
        kernel_skinny = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
        dilated_skinny = cv2.dilate(dilated, kernel_skinny)
        # show img with dilated edges on kanny edge detection
        #im2b = fig_edges.add_subplot(133)  # right
        #im2b.imshow(dilated_skinny)
        #im2b.set_title("dilated_skinny")
        cv2.imwrite(path + "dilated_skinny.jpg", dilated_skinny)
        # add success for this process to the config
        self.config_file.dilated_skinny_output_config_file(path + "dilated_skinny.jpg")
        new_contours_skinny, hierarchy_skinny = cv2.findContours(dilated_skinny.copy(),
                                                                 cv2.RETR_EXTERNAL,
                                                                 cv2.CHAIN_APPROX_NONE)
        contours_skinny = sorted(new_contours_skinny, key=cv2.contourArea, reverse=True)[:10]
        contour_skinny = contours_skinny[0]

        # draw contours
        gray_for_drawing_1 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.drawContours(gray_for_drawing_1, contour, -1, (0, 0, 255), 2)
        cv2.imwrite(path + ".contours.jpg", gray_for_drawing_1)
        # add success for this process to the config
        self.config_file.contours_output_config_file(path + ".contours.jpg")
        # draw contours skinny
        gray_for_drawing_1_skinny = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.drawContours(gray_for_drawing_1_skinny, contour_skinny, -1, (0, 0, 255), 2)
        cv2.imwrite(path + ".contours_skinny.jpg", gray_for_drawing_1_skinny)
        # add success for this process to the config
        self.config_file.contours_skinny_output_config_file(path + ".contours_skinny.jpg")
        # find centroid
        m = cv2.moments(contour)
        centroid_x = int(m['m10'] / m['m00'])
        centroid_y = int(m['m01'] / m['m00'])
        centroid = (centroid_x, centroid_y)
        gray_for_drawing_2 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        # print "output centroid image: ", path + ".centroid.jpg"

        cv2.circle(gray_for_drawing_2, centroid, 3, thickness=-4, color=(0, 0, 255))
        cv2.imwrite(path + ".centroid.jpg", gray_for_drawing_2)
        # add success for this process to the config
        self.config_file.centroid_output_config_file(path + ".centroid.jpg")
        # find centroid skinny
        m = cv2.moments(contour_skinny)
        centroid_x = int(m['m10'] / m['m00'])
        centroid_y = int(m['m01'] / m['m00'])
        centroid_skinny = (centroid_x, centroid_y)
        gray_for_drawing_2_skinny = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.circle(gray_for_drawing_2_skinny, centroid_skinny, 3, thickness=-4, color=(0, 0, 255))
        cv2.imwrite(path + ".centroid_skinny.jpg", gray_for_drawing_2_skinny)
        # add success for this process to the config
        self.config_file.centroid_skinny_output_config_file(path + ".centroid_skinny.jpg")
        # find a bounding rectangle for the image
        x, y, w, h = cv2.boundingRect(contour)
        cv2.rectangle(image, (x, y), (x + w, y + h), (0, 0, 0), 1)
        # cv2.imwrite(path + ".contours.centroid.epts.vl.br.jpg", image)
        # creating lamina length
        offset = -10
        # top line
        p1 = (x + offset, y)
        p2 = (x + w, y)
        gray_for_drawing_3 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.line(gray_for_drawing_3, p1, p2, (0, 0, 255), 1)
        # bottom line
        p1 = (x + offset, y + h)
        p2 = (x + w, y + h)
        cv2.line(gray_for_drawing_3, p1, p2, (0, 0, 255), 1)
        # vert line
        p1 = (x, y)
        p2 = (x, y + h)
        cv2.arrowedLine(gray_for_drawing_3, p1, p2, (0, 0, 255), 1, tipLength=0.05)
        cv2.arrowedLine(gray_for_drawing_3, p2, p1, (0, 0, 255), 1, tipLength=0.05)
        cv2.imwrite(path + ".length.jpg", gray_for_drawing_3)
        # add success for this process to the config
        self.config_file.length_output_config_file(path + ".length.jpg")
        # creating lamina width
        offset = 10
        # top line
        p1 = (x, y)
        p2 = (x, y + h + offset)
        gray_for_drawing_4 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.line(gray_for_drawing_4, p1, p2, (0, 0, 255), 1)
        # bottom line
        p1 = (x + w, y)
        p2 = (x + w, y + h + offset)
        cv2.line(gray_for_drawing_4, p1, p2, (0, 0, 255), 1)
        # vert line
        p1 = (x, y + h + 2)
        p2 = (x + w, y + h + 2)
        cv2.arrowedLine(gray_for_drawing_4, p1, p2, (0, 0, 255), 1, tipLength=0.05)
        cv2.arrowedLine(gray_for_drawing_4, p2, p1, (0, 0, 255), 1, tipLength=0.05)
        cv2.imwrite(path + ".width.jpg", gray_for_drawing_4)
        # add success for this process to the config
        self.config_file.width_output_config_file(path + ".width.jpg")
        # find area
        area = cv2.contourArea(contour)
        gray_for_drawing_5 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        # contour_in_a_list_for_param = []
        # contour_in_a_list_for_param.append(contour)
        contour_in_a_list_for_param = [contour]
        cv2.fillPoly(gray_for_drawing_5, contour_in_a_list_for_param, (0, 0, 255))
        cv2.imwrite(path + ".area.jpg", gray_for_drawing_5)
        # add success for this process to the config
        self.config_file.area_output_config_file(path + ".area.jpg")
        # find perimeter
        perimeter = cv2.arcLength(contour, True)

        # aspect_ratio = length/width
        aspect_ratio = float(float(h) / float(w))

        # creating lamina length for aspect ratio
        # vertical line
        p1 = (x + (w / 2), y)
        p2 = (x + (w / 2), y + h)
        gray_for_drawing_6 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        PeaUtils.drawline(gray_for_drawing_6, p1, p2, (0, 0, 255))
        # horizontal line
        p1 = (x, y + (h / 2))
        p2 = (x + w, y + (h / 2))
        PeaUtils.drawline(gray_for_drawing_6, p1, p2, (0, 0, 255))
        cv2.imwrite(path + ".aspectratio.jpg", gray_for_drawing_6)
        # add success for this process to the config
        self.config_file.aspectratio_output_config_file(path + ".aspectratio.jpg")
        # roundness = 4piA/(P*P)
        roundness = (4 * math.pi * area) / (perimeter * perimeter)
        gray_for_drawing_7 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.circle(gray_for_drawing_7, centroid, (w / 2), (0, 0, 255), 1)
        cv2.imwrite(path + ".roundness.jpg", gray_for_drawing_7)
        # add success for this process to the config
        self.config_file.roundness_output_config_file(path + ".roundness.jpg")
        # compactness = (Perimeter * Perimeter) / Area
        compactness = (perimeter * perimeter) / area

        gray_for_drawing_8 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.drawContours(gray_for_drawing_8, contour, -1, (255, 0, 0), 5)
        # contour_in_a_list_for_param = []
        # contour_in_a_list_for_param.append(contour)
        contour_in_a_list_for_param = [contour]
        cv2.fillPoly(gray_for_drawing_8, contour_in_a_list_for_param, (0, 0, 255))
        cv2.imwrite(path + ".compactness.jpg", gray_for_drawing_8)
        # add success for this process to the config
        self.config_file.compactness_output_config_file(path + ".compactness.jpg")
        # rectangularity = Area / (length*width)
        rectangularity = area / (h * w)

        x_coord, y_coord, w_coord, h_coord = cv2.boundingRect(contour)
        gray_for_drawing_9 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.rectangle(gray_for_drawing_9, (x_coord, y_coord), (x_coord + w_coord, y_coord + h_coord), (0, 0, 255), 1)
        cv2.imwrite(path + ".rectangularity.jpg", gray_for_drawing_9)
        # add success for this process to the config
        self.config_file.rectangularity_output_config_file(path + ".rectangularity.jpg")
        # perimeter ratio of major axis length
        perimeter_ratio_length = perimeter / h

        gray_for_drawing_10 = gray_for_drawing_3
        cv2.drawContours(gray_for_drawing_10, contour, -1, (255, 0, 0), 2)
        cv2.imwrite(path + ".perimeterratiolength.jpg", gray_for_drawing_10)
        # add success for this process to the config
        self.config_file.perimeterratiolength_output_config_file(path + ".perimeterratiolength.jpg")
        # perimeter ratio of major /minor length Plw = P / (l+w)
        plw = perimeter / (h + w)

        gray_for_drawing_11 = gray_for_drawing_6
        cv2.drawContours(gray_for_drawing_11, contour, -1, (255, 0, 0), 2)
        cv2.imwrite(path + ".perimeterratiolengthwidth.jpg", gray_for_drawing_11)
        # add success for this process to the config
        self.config_file.perimeterratiolengthwidth_output_config_file(path + ".perimeterratiolengthwidth.jpg")
        # convex hull
        hull = cv2.convexHull(contour)
        convexhull_area = cv2.contourArea(hull)
        convexhull_perimeter = cv2.arcLength(hull, True)

        # perimeter convexity = Pch / P
        perimeter_convexity = convexhull_perimeter / perimeter

        gray_for_drawing_12 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        pts = np.array(hull, np.int32)
        pts = pts.reshape((-1, 1, 2))
        cv2.polylines(gray_for_drawing_12, [pts], True, (255, 0, 0))
        cv2.imwrite(path + ".perimeterconvexity.jpg", gray_for_drawing_12)
        # add success for this process to the config
        self.config_file.perimeterconvexity_output_config_file(path + ".perimeterconvexity.jpg")
        # area convexity = (convex hull area -area )/ area
        area_convexity = (convexhull_area - area) / area

        gray_for_drawing_13 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        pts = np.array(hull, np.int32)
        pts = pts.reshape((-1, 1, 2))
        cv2.fillPoly(gray_for_drawing_13, [pts], (255, 0, 0))
        # contour_in_a_list_for_param = []
        # contour_in_a_list_for_param.append(contour)
        contour_in_a_list_for_param = [contour]
        cv2.fillPoly(gray_for_drawing_13, contour_in_a_list_for_param, (0, 0, 255))
        cv2.imwrite(path + ".areaconvexity.jpg", gray_for_drawing_13)
        # add success for this process to the config
        self.config_file.areaconvexity_output_config_file(path + ".areaconvexity.jpg")
        # area ratio of convexity = A / convex hull area
        area_ratio_of_convexity = area / convexhull_area

        cv2.imwrite(path + ".arearatioconvexity.jpg", gray_for_drawing_13)
        # add success for this process to the config
        self.config_file.arearatioconvexity_output_config_file(path + ".arearatioconvexity.jpg")
        # equivalent diameter = sqrt(4A/pi)
        eqd = math.sqrt((4 * area) / math.pi)

        gray_for_drawing_14 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        cv2.circle(gray_for_drawing_14, centroid, int(eqd / 2), (0, 0, 255), 1)
        cv2.imwrite(path + ".equivalentdiameter.jpg", gray_for_drawing_14)
        # add success for this process to the config
        self.config_file.equivalentdiameter_output_config_file(path + ".equivalentdiameter.jpg")
        ################
        # TOOTH FINDING
        ################
        print "Making tooth."
        # contour_image_path_skinny = path + ".contours_skinny.jpg"
        contour_image_path = path + ".contours.jpg"
        # centroid_image_path_skinny=path + ".centroid_skinny.jpg"
        centroid_image_path = path + ".centroid.jpg"
        # TODO: hangs at number_teeth
        number_teeth = self.find_teeth(centroid_image_path, contour_image_path, centroid, contour_skinny)
        # TODO: ratio for serration
        print "number of teeth: ", number_teeth
        # TODO: look at tooth whole tooth nothing but the tooth for additional metrics to add in!

        if number_teeth > 3:
            serrated = True
        else:
            serrated = False

        #########################################
        # Block for constructing info string.
        #########################################
        final_l = ""
        for lab in self.image_label.imageLabelList:
            final_l += lab + " "
        pix_per_cm = self.image_ruler.pixels_per_cm
        length_in_cm = float(h) / float(pix_per_cm)
        width_in_cm = float(w) / float(pix_per_cm)

        # pix_per_mm=((pix_per_cm/2)/10)
        # square_pix=     float(1)/float((pix_per_mm*pix_per_mm))
        # area_sqr_mm =   float(area * square_pix)
        # area_in_cm_sq = float(area_sqr_mm)/float(10)
        # area_in_cm_sq = float(area_in_cm_sq)
        one_over_pix_per_cm = float(1.0 / pix_per_cm)
        area_in_cm_sq = area * one_over_pix_per_cm * one_over_pix_per_cm
        perimeter_in_cm = float(perimeter) / float(pix_per_cm)
        # eqd in cm
        eqd_cm = eqd / pix_per_cm

        a1 = str(final_l)
        a2 = str(pix_per_cm)
        a3 = str(h)
        a4 = str(length_in_cm)
        a5 = str(w)
        a6 = str(width_in_cm)
        a7 = str(area)
        a7a = str(area_in_cm_sq)
        a8 = str(perimeter)
        a9 = str(perimeter_in_cm)
        a10 = str(centroid)
        a11 = str(aspect_ratio)
        a12 = str(roundness)
        a13 = str(compactness)
        a14 = str(rectangularity)
        a15 = str(perimeter_ratio_length)
        a16 = str(plw)
        a17 = str(perimeter_convexity)
        a18 = str(area_convexity)
        a19 = str(area_ratio_of_convexity)
        a20 = str(eqd)
        a20a = str(eqd_cm)
        a21 = str(serrated)
        a22 = str(number_teeth)
        self.csv_string = a1 + "\t" + a2 + "\t" + a3 + "\t" + a4 + "\t" + a5 + "\t" + a6 + "\t" + a7 + "\t" + a7a + "\t" + a8 + "\t" + a9 + "\t" + a10 + "\t" + a11 + "\t" + a12 + "\t" + a13 + "\t" + a14 + "\t" + a15 + "\t" + a16 + "\t" + a17 + "\t" + a18 + "\t" + a19 + "\t" + a20 + "\t" + a20a + "\t" + a21 + "\t" + a22
        #print "pea image = self.csv_string"
        print self.csv_string
        print "Making word document."
        document = docx.Document('../resources/worddoc/landscape-template.docx')
        style = document.styles['Normal']
        font = style.font
        font.name = 'Baskerville'
        # font.size
        document.add_paragraph('MktStall v1.01')
        document.add_paragraph('Morphological Descriptors\tTO:0000829/TO:0000748')
        document.add_paragraph(
            'A cardinal organ part morphology trait (TO:0000754) which is a quality of a leaf lamina ('
            'PO:0020039)/phyllome morphology trait (TO:0000747) which is quality of a leaf (PO:0025034)')
        paragraph = document.add_paragraph('')
        paragraph.add_run(a1).bold = True
        paragraph.add_run('\n')
        table = document.add_table(rows=1, cols=7)
        table.add_row().cells
        heading_cells = table.rows[0].cells
        heading_cells[0].text = 'Descriptor'
        heading_cells[1].text = 'Ontology Identifier'
        heading_cells[2].text = 'Ontology Description'
        heading_cells[3].text = 'MktStall Explanation'
        heading_cells[4].text = 'Picture'
        heading_cells[5].text = 'Pixel'
        heading_cells[6].text = 'Cm'

        morphometric_descriptors = ["Major axis length, L", "Minor axis length, W", "Area, A", "Perimeter, P",
                                    "Leaf Margin Serrated, Ser", "Leaf Lamina Tooth, Too"]
        ontology_identifiers = ["TO:0002690", "TO:0002720", "TO:0000827", "TO:0000616", "TO:0006063", "PO:0025518"]
        ont_desc = [
            "A leaf lamina morphology trait (TO:0000829) which is the length of the leaf laimina (PO:0020039). Refer "
            "to length (PATO:0000122): A 1-D extent quality which is equal to the distance between two points. "
            "Usually measured in centimeters.",
            "Refer to width (PATO:0000921): A 1-D extent quality which is equal to the distance from one side of an "
            "object to another side which is opposite. in centimeters A leaf lamina morphology trait (TO:0000829) "
            "which is the width of the leaf lamina (PO:0020039).",
            "A leaf lamina morphology trait (TO:0000829) which is the area of a leaf lamina (PO:0020039).",
            "It is a measure of the total length covered by the leaf margin.",
            "A leaf margin morphology trait (TO:0002635) which is a notched or grooved leaf margin (PO:0020128). "
            "Refer to serrated (PATO:0001206): A shape quality inhering in a bearer by virtue of having sharp "
            "straight-edged teeth pointing to the apex.",
            "A phyllome lamina tooth (PO:0025515) that is an angular projection on a leaf margin (PO:0020128) for "
            "which the corresponding leaf sinus (PO:0025384) extends less than one quarter of the distance to the "
            "center of the long axis of the leaf lamina (PO:0020039)."]
        mktstall_explanation = ["Distance between the base and the tip of the leaf",
                                "Maximum width of the leaf that is perpendicular to the major axis",
                                "Number of pixels found within the contour of the leaf",
                                "Distance of the number of pixels found in the leaf contour",
                                "Presence of leaf teeth on the margin",
                                "Number of leaf teeth found of the margin"]

        path_pic_length = path + ".length.jpg"
        path_pic_width = path + ".width.jpg"
        path_pic_area = path + ".area.jpg"
        path_pic_contours = path + ".contours.jpg"
        path_pic_teeth = centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_lobeF.jpg"


        list_of_picture_paths = [path_pic_length,
                                 path_pic_width,
                                 path_pic_area,
                                 path_pic_contours,
                                 path_pic_contours,
                                 path_pic_teeth]
        list_in_pixels = [a3, a5, a7, a8, a21, a22]
        list_in_cm = [a4, a6, "", a9, a21, a22]
        ################################################

        for x in range(0, 6):
            cells = table.add_row().cells
            cells[0].text = morphometric_descriptors[x]
            cells[1].text = ontology_identifiers[x]
            cells[2].text = ont_desc[x]
            cells[3].text = mktstall_explanation[x]
            paragraph = cells[4].paragraphs[0]
            run = paragraph.add_run()
            if os.path.isfile(list_of_picture_paths[x]):
                img_path = self.resize_image_pc(list_of_picture_paths[x], 0.5)
                img = Image.open(img_path)
                dimensions = img.size
                w = dimensions[0] * 10000
                h = dimensions[1] * 10000
                run.add_picture(list_of_picture_paths[x], width=w, height=h)
                cells[5].text = list_in_pixels[x]
                cells[6].text = list_in_cm[x]
                x += 1
        label_into_joined_name = ",".join(a1)

        document.add_page_break()
        document.add_paragraph('MktStall v1.01.')
        document.add_paragraph('Shape Descriptors')
        document.add_paragraph('TO:0000492')
        document.add_paragraph(
            'A leaf morphology trait (TO:0000748) which is the variation in shapes and forms of a leaf (PO:0025034). '
            'Refer to shape (PATO:0000052): A morphological quality inhering in a bearer by virtue of the bearer\'s '
            'ratios of distances between its features (points, edges, surfaces and also holes etc). If you are '
            'annotating to this term, please add an additional annotation to vascular leaf morphology (TO:0000419) or '
            'non-vascular leaf morphology trait (TO:0000925), depending on the species.')
        document.add_paragraph(a1)

        shapetable = document.add_table(rows=1, cols=7)
        row_cells = table.add_row().cells
        shape_heading_cells = table.rows[0].cells
        shape_heading_cells[0].text = 'Descriptor'
        shape_heading_cells[1].text = 'Ontology Identifier'
        shape_heading_cells[2].text = 'Ontology Description'
        shape_heading_cells[3].text = 'MktStall Explanation'
        shape_heading_cells[4].text = 'Picture'
        shape_heading_cells[5].text = 'Pixel'
        shape_heading_cells[6].text = 'Cm'
        shape_desc = ["Centroid",
                      "Aspect Ratio, AR",
                      "Roundness, R",
                      "Compactness, C",
                      "Rectangularity, N",
                      "Perimeter ratio of Major Axis Length, PL",
                      "Perimeter ratio of Major axis length and Minor axis length, PLW",
                      "Perimeter convexity, PC",
                      "Area convexity (aka Entirety), AC1",
                      "Area ratio of convexity (aka Solidity), AC2",
                      "Equivalent Diameter, EqD"]

        shape_ontology_identifiers = ["", "TO:0000542", "", "", "", "", "", "", "", "", ""]

        shape_ont_desc = ["",
                          "A leaf size (TO:0002637) which is the ratio of leaf length to its width. If you are "
                          "annotating to this term, please add an additional annotation to vascular leaf morphology ("
                          "TO:0000419) or non-vascular leaf morphology trait (TO:0000925), depending on the species.",
                          "", "", "", "", "", "", "", "", ""]
        shape_mktstall_explanation = ["The coordinates of the leaf center",
                                      "The ratio of major and minor axis lengths",
                                      "The difference between the leaf and a circle",
                                      "Ratio of perimeter squared over area",
                                      "How rectangular the leaf is",
                                      "Ratio of perimeter over length",
                                      "Ratio of perimeter over length plus width",
                                      "Ratio of convex perimeter over leaf perimeter",
                                      "Difference of convex hull area and area over area (normalised)",
                                      "Ratio of leaf area and convex hull area",
                                      "The diameter of a circle that has the same area as the leaf"]

        # TESTER
        # string generator http://www.mathmlcentral.com/Tools/ToMathML.jsp
        # roundness_mathml_string =
        # '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mn>1</mn><mn>2</mn></mfrac></math>'
        mathml_string_list = []

        centroid_xy_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mi>(x,y)</mi></math>'
        aspect_ratio_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"> <mfrac><mi>L</mi><mi>W</mi></mfrac></math>'
        roundness_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mrow><mn>4</mn><mo>&#8290;</mo><mi>&#960;</mi><mo>&#8290;</mo><mi>A</mi></mrow><msup><mi>P</mi><mn>2</mn></msup></mfrac></math>'
        compactness_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><msup><mi>P</mi><mn>2</mn></msup><mi>A</mi></mfrac></math>'
        rectangularity_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>A</mi><mrow><mi>L</mi><mo>&#8290;</mo><mi>W</mi></mrow></mfrac></math>'
        perimeter_ratio_of_major_axis_length_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>P</mi><mi>L</mi></mfrac></math>'
        perimeter_ratio_of_major_axis_length_width_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>P</mi><mrow><mi>L</mi><mo>+</mo><mi>W</mi></mrow></mfrac></math>'
        perimeter_convexity_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>Pch</mi><mi>P</mi></mfrac></math>'
        area_convexity_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>Ach</mi><mi>A</mi></mfrac></math>'
        area_ratio_convexity_mathml_string = '<math xmlns="http://www.w3.org/1998/Math/MathML"><mfrac><mi>A</mi><mi>Ach</mi></mfrac></math>'
        equivalent_diameter_mathml_string = '<math xmlns ="http://www.w3.org/1998/Math/MathML"><mrow><mi>sqrt</mi><mo>&#8289;</mo><mo>(</mo><mfrac><mrow><mn>4</mn><mo>&#8290;</mo><mi>A</mi></mrow><mi>pi</mi></mfrac><mo>)</mo></mrow></math>'

        mathml_string_list.append(centroid_xy_string)
        mathml_string_list.append(aspect_ratio_mathml_string)
        mathml_string_list.append(roundness_mathml_string)
        mathml_string_list.append(compactness_mathml_string)
        mathml_string_list.append(rectangularity_mathml_string)
        mathml_string_list.append(perimeter_ratio_of_major_axis_length_mathml_string)
        mathml_string_list.append(perimeter_ratio_of_major_axis_length_width_mathml_string)
        mathml_string_list.append(perimeter_convexity_mathml_string)
        mathml_string_list.append(area_convexity_mathml_string)
        mathml_string_list.append(area_ratio_convexity_mathml_string)
        mathml_string_list.append(equivalent_diameter_mathml_string)
        formula = []


        for s in mathml_string_list:
            tree = etree.fromstring(s)
            xslt = etree.parse("../resources/worddoc/MML2OMML.XSL")
            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            paragraph = document.add_paragraph()
            paragraph._element.append(new_dom.getroot())
            formula.append(new_dom)
            # add to formula
        # end of tester

        path_pic_roundness = path + ".roundness.jpg"
        path_pic_compactness = path + ".compactness.jpg"
        path_pic_rectangularity = path + ".rectangularity.jpg"
        path_pic_perimeter_ratio_length = path + ".perimeterratiolength.jpg"
        path_pic_perimeter_ratio_length_width = path + ".perimeterratiolengthwidth.jpg"
        path_pic_perimeter_convexity = path + ".perimeterconvexity.jpg"
        path_pic_area_convexity = path + ".areaconvexity.jpg"
        path_pic_area_ratio_convexity = path + ".arearatioconvexity.jpg"
        path_pic_equivalent_diameter = path + ".equivalentdiameter.jpg"
        path_pic_centroid = path + ".centroid.jpg"
        path_pic_aspect_ratio = path + ".aspectratio.jpg"
        shape_list_of_picture_paths = [path_pic_centroid,
                                      path_pic_aspect_ratio,
                                      path_pic_roundness,
                                      path_pic_compactness,
                                      path_pic_rectangularity,
                                      path_pic_perimeter_ratio_length,
                                      path_pic_perimeter_ratio_length_width,
                                      path_pic_perimeter_convexity,
                                      path_pic_area_convexity,
                                      path_pic_area_ratio_convexity,
                                      path_pic_equivalent_diameter]
        shape_list_in_pixels = [a10, a11, a12, a13, a14, a15, a16, a17, a18, a19, a20]

        for x in range(0, 11):
            cells = shapetable.add_row().cells
            cells[0].text = shape_desc[x]
            cells[1].text = shape_ontology_identifiers[x]
            cells[2].text = shape_ont_desc[x]
            cells[3].text = shape_mktstall_explanation[x]
            if formula[x] != "":
                cells[4].paragraphs[0]._element.append(formula[x].getroot())
            else:
                cells[4].text = formula[x]
            paragraph = cells[5].paragraphs[0]
            run = paragraph.add_run()
            img_path = self.resize_image_pc(shape_list_of_picture_paths[x], 0.5)
            img = Image.open(img_path)
            dimensions = img.size
            w = dimensions[0] * 10000
            h = dimensions[1] * 10000
            run.add_picture(shape_list_of_picture_paths[x], width=w, height=h)
            cells[6].text = shape_list_in_pixels[x]

        joined_final_l = final_l.replace(" ", "_")
        self.joined_final_l = joined_final_l
        self.config_file.label_string_output_config_file(self.joined_final_l, self.image_label.imageAcceptable)
        document.save(Parameters.results + '/worddoc/' + joined_final_l + "demo.docx")
        ################################################
        # html conversion
        ################################################
        #
        # dictionary morphometric + shape
        # accession name, picture pixel cm

        # STEP ONE
        # make an x by y matrix here where each row is a phenotype, and each coloumn is a measure/image/descript etc.

        # morphological 7 across 7 down
        m_across = 7
        m_down = 7
        m_matrix = [""] * m_across
        for i in range(m_across):
            m_matrix[i] = [""] * m_down

        header = ["Descriptor", "Ontology Identifier", "Ontology Description", "MktStall Explanation", "Picture",
                  "Pixel", "Cm"]
        m_pheno_1 = ["Major axis length, L",
                     "TO:0002690",
                     "A leaf lamina morphology trait (TO:0000829) which is the length of the leaf laimina ("
                     "PO:0020039). Refer to length (PATO:0000122): A 1-D extent quality which is equal to the "
                     "distance between two points. Usually measured in centimeters.",
                     "Distance between the base and the tip of the leaf",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_length,
                     a3,
                     a4]
        m_pheno_2 = ["Minor axis length, W",
                     "TO:0002720",
                     "Refer to width (PATO:0000921): A 1-D extent quality which is equal to the distance from one "
                     "side of an object to another side which is opposite. in centimeters A leaf lamina morphology "
                     "trait (TO:0000829) which is the width of the leaf lamina (PO:0020039).",
                     "Maximum width of the leaf that is perpendicular to the major axis",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_width,
                     a5,
                     a6]
        m_pheno_3 = ["Area, A",
                     "TO:0000827",
                     "A leaf lamina morphology trait (TO:0000829) which is the area of a leaf lamina (PO:0020039).",
                     "Number of pixels found within the contour of the leaf",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_area,
                     a7,
                     a7a]

        m_pheno_4 = ["Perimeter, P",
                     "TO:0000616",
                     "It is a measure of the total length covered by the leaf margin.",
                     "Distance of the number of pixels found in the leaf contour",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_contours,
                     a8,
                     a9]

        m_pheno_5 = ["Leaf Margin Serrated, Ser",
                     "TO:0006063",
                     "A leaf margin morphology trait (TO:0002635) which is a notched or grooved leaf margin ("
                     "PO:0020128). Refer to serrated (PATO:0001206): A shape quality inhering in a bearer by virtue "
                     "of having sharp straight-edged teeth pointing to the apex.",
                     "Presence of leaf teeth on the margin",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_contours,
                     a21,
                     a21]
        m_pheno_6 = ["Leaf Lamina Tooth, Too",
                     "PO:0025518",
                     "(PO:0025515) that is an angular projection on a leaf margin (PO:0020128) for which the "
                     "corresponding leaf sinus (PO:0025384) extends less than one quarter of the distance to the "
                     "center of the long axis of the leaf lamina (PO:0020039).",
                     "Number of leaf teeth found of the margin",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_teeth,
                     a22,
                     a22]

        m_matrix[0] = header
        m_matrix[1] = m_pheno_1
        m_matrix[2] = m_pheno_2
        m_matrix[3] = m_pheno_3
        m_matrix[4] = m_pheno_4
        m_matrix[5] = m_pheno_5
        m_matrix[6] = m_pheno_6

        # print "here is your matrix", m_matrix

        # Step two
        # make a new matrix where each row is a string containing the html code + the row data from step one.
        html_table = ""
        html_head_start = "<thead>\n<tr>\n"
        html_head_content = ""
        for i in m_matrix[0]:
            html_head_content = html_head_content + "<th>\n" + i + "\n</th>\n"
        html_head_end = "</tr>\n</thead>\n"

        html_body_start = "<tbody>\n<tr>\n"
        html_body_content = ""
        phenotypes = m_matrix[1:(len(m_matrix))]
        for row in phenotypes:
            html_body_content = html_body_content + "<tr>\n"
            for data in row:
                html_body_content = html_body_content + "<td>\n" + data + "</td>\n"
            html_body_content = html_body_content + "</tr>\n"
        html_body_end = "</tr>\n</tbody>\n"

        html_table = html_head_start +\
                     html_head_content +\
                     html_head_end +\
                     html_body_start +\
                     html_body_content +\
                     html_body_end

        # shape 7 across 7 down
        s_across = 7  # this many items
        s_down = 12  # this many lists
        s_matrix = [[0 for x in range(s_across)] for y in range(s_down)]


        header = ["Descriptor",
                  "Ontology Identifier",
                  "Ontology Description",
                  "MktStall Explanation",
                  "Formula <a href=\"https://link.springer.com/article/10.1007/s11831-016-9206-z\">(Waldchen <i>et al</i>., 2016)</a>",
                  "Picture",
                  "Value"]

        s_pheno_1 = ["Centroid",
                     "",
                     "",
                     "The coordinates of the leaf center",
                     "$${(x,y)}$$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_centroid,
                     a10]
        s_pheno_2 = ["Aspect Ratio, AR",
                     "TO:0000542",
                     "A leaf size (TO:0002637) which is the ratio of leaf length to its width. If you are annotating "
                     "to this term, please add an additional annotation to vascular leaf morphology (TO:0000419) or "
                     "non-vascular leaf morphology trait (TO:0000925), depending on the species.",
                     "The ratio of major and minor axis lengths",
                     "$$ {L \over W} $$ ",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_aspect_ratio,
                     a11]
        s_pheno_3 = ["Roundness, R",
                     "",
                     "",
                     "The difference between the leaf and a circle",
                     "$${4 \pi A \over P^2}$$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_roundness,
                     a12]

        s_pheno_4 = ["Compactness, C",
                     "",
                     "",
                     "Ratio of perimeter squared over area",
                     "$${P^2 \over A}$$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_compactness,
                     a13]

        s_pheno_5 = ["Rectangularity, N",
                     "",
                     "",
                     "Describes how rectangular the leaf is.",
                     "$$ {A \over LW} $$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_rectangularity,
                     a14]

        s_pheno_6 = ["Perimeter ratio of Major Axis Length, PL",
                     "",
                     "",
                     "Ratio of perimeter over major axis length",
                     "$$ {P \over L }$$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_perimeter_ratio_length,
                     a15]

        s_pheno_7 = ["Perimeter Ratio of Major and Minor Axis Length, PLW",
                     "",
                     "",
                     "Perimeter ratio of major and minor axis length",
                     "$$ {P \over {L + W}} $$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_perimeter_ratio_length_width,
                     a16]

        s_pheno_8 = ["Perimeter Convexity, PC",
                     "",
                     "",
                     "Perimeter ratio of convex hull perimeter and perimeter",
                     "$$ {P_{ch} \over P} $$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_perimeter_convexity,
                     a17]

        s_pheno_9 = ["Area Convexity (aka Entirety), AC1",
                     "",
                     "",
                     "Ratio of the difference between area of the convex hull and the area.",
                     "$$ {A_{ch} \over A} $$",
                     "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_area_convexity,
                     a18]

        s_pheno_10 = ["Area Ratio of Convexity (aka Solidity), AC2",
                      "",
                      "",
                      "Ratio of area over area of the convexhull.",
                      "$$ {A \over A_{ch}} $$",
                      "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_area_ratio_convexity,
                      a19]

        s_pheno_11 = ["Equivalent Diameter, EqD",
                      "",
                      "",
                      "The diameter of a circle that would have the same area as the area of the leaf.",
                      "$$ {\sqrt{4A \over \pi}} $$ ",
                      "<img src = \"%s\" class =\"img-fluid\" alt=\"Responsive image\" >" % path_pic_equivalent_diameter,
                      a20]

        s_matrix[0] = header
        s_matrix[1] = s_pheno_1
        s_matrix[2] = s_pheno_2
        s_matrix[3] = s_pheno_3
        s_matrix[4] = s_pheno_4
        s_matrix[5] = s_pheno_5
        s_matrix[6] = s_pheno_6
        s_matrix[7] = s_pheno_7
        s_matrix[8] = s_pheno_8
        s_matrix[9] = s_pheno_9
        s_matrix[10] = s_pheno_10
        s_matrix[11] = s_pheno_11

        # print "heres your matrix", m_matrix

        # Step two
        # make a new matrix where each row is a string containing the html code + the row data from step one.
        # s_html_table = ""
        s_html_head_start = "<thead>\n<tr>\n"
        s_html_head_content = ""
        for i in s_matrix[0]:
            s_html_head_content = s_html_head_content + "<th>\n" + i + "\n</th>\n"
        s_html_head_end = "</tr>\n</thead>\n"

        s_html_body_start = "<tbody>\n<tr>\n"
        s_html_body_content = ""
        s_phenotypes = s_matrix[1:(len(s_matrix))]
        for row in s_phenotypes:
            s_html_body_content = s_html_body_content + "<tr>\n"
            for data in row:
                s_html_body_content = s_html_body_content + "<td>\n" + str(data) + "</td>\n"
            s_html_body_content = s_html_body_content + "</tr>\n"
        s_html_body_end = "</tr>\n</tbody>\n"

        s_html_table = s_html_head_start + s_html_head_content + s_html_head_end + s_html_body_start + s_html_body_content + s_html_body_end
        # print s_html_table
        ##########################


        # step three
        # substitute/ADD the lines in myHTMLtESTDoc which the lice that you have made in step two
        my_html_test_doc = []
        """Morphological descriptors"""
        f = open("../resources/html/index.html", "r")
        my_html_test_doc = f.read().split('\n')
        # print myHtmlTestDoc
        # substitute html_table for "THIS IS TABLE"
        html_index = my_html_test_doc.index("THIS IS TABLE")
        my_html_test_doc[html_index] = html_table

        html_index = my_html_test_doc.index("<link href=\"bootstrap.min.css\" rel=\"stylesheet\">")
        bootstrap_string = "<link href=\""+Parameters.resources+"/css/bootstrap.min.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = bootstrap_string

        html_index = my_html_test_doc.index("<link href=\"dashboard.css\" rel=\"stylesheet\">")
        dashboard_string = "<link href=\""+ Parameters.resources+"/css/dashboard.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = dashboard_string

        f.close()
        html_file_name_morphological = Parameters.results + \
                                       "/html/html_morphological_descriptors" + \
                                       "_" + \
                                       joined_final_l + \
                                       ".html"
        new_html = open(html_file_name_morphological, 'w')
        str_for_html = "".join(my_html_test_doc)
        new_html.write(str_for_html)
        # step four
        # output the html to file and open in a browser.
        filename = 'file:///' + html_file_name_morphological

        # add success for this process to the config
        self.config_file.html_morphological_output_config_file(html_file_name_morphological)
        """Shape descriptors"""
        f = open("../resources/html/index.html", "r")
        my_html_test_doc = f.read().split('\n')
        html_index = my_html_test_doc.index("THIS IS TABLE")
        my_html_test_doc[html_index] = s_html_table

        html_index = my_html_test_doc.index("<link href=\"bootstrap.min.css\" rel=\"stylesheet\">")
        bootstrap_string = "<link href=\"" + Parameters.resources + "/css/bootstrap.min.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = bootstrap_string

        html_index = my_html_test_doc.index("<link href=\"dashboard.css\" rel=\"stylesheet\">")
        dashboard_string = "<link href=\"" + Parameters.resources + "/css/dashboard.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = dashboard_string

        f.close()
        html_file_name_shape = Parameters.results + \
                                       "/html/html_shape_descriptors" + \
                                       "_" + \
                                       joined_final_l + \
                                       ".html"
        new_html = open(html_file_name_shape, 'w')
        str_for_html = "".join(my_html_test_doc)
        new_html.write(str_for_html)
        # step four
        # output the html to file and open in a browser.
        filename = 'file:///' + html_file_name_shape

        # add success for this process to the config
        self.config_file.html_shape_output_config_file(html_file_name_shape)
        """Leaf Margin Chart"""
        f = open("../resources/html/index_chart.html", "r")
        my_html_test_doc = f.read().split('\n')
        html_index = my_html_test_doc.index("data: [15339, 21345, 18483, 24003, 23489, 24092, 12034],")
        results_str = "data: %s," % self.distances
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index(
           "labels: [\"Sunday\", \"Monday\", \"Tuesday\", \"Wednesday\", \"Thursday\", \"Friday\", \"Saturday\"],")
        results_str = "labels: %s," % self.new_position_index_list
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [60, 60, 60, 60, 60, 60, 60],")
        self.mean_dist_from_centroid = self.mean_dist_from_centroid * len(self.new_position_index_list)
        results_str = "data: %s," % self.mean_dist_from_centroid
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [70, 70, 70, 70, 70, 70, 70],")
        self.onestd_multiplierabove = self.onestd_multiplierabove * len(self.new_position_index_list)
        results_str = "data: %s," % self.onestd_multiplierabove
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [71, 71, 71, 71, 71, 71, 71],")
        self.onestd_multiplierbelow = self.onestd_multiplierbelow * len(self.new_position_index_list)
        results_str = "data: %s," % self.onestd_multiplierbelow
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [80, 80, 80, 80, 80, 80, 80],")
        self.one_std_above = self.one_std_above * len(self.new_position_index_list)
        results_str = "data: %s," % (self.one_std_above)
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [81, 81, 81, 81, 81, 81, 81],")
        self.one_std_below = self.one_std_below * len(self.new_position_index_list)
        results_str = "data: %s," % (self.one_std_below)
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [90, 90, 90, 90, 90, 90, 90],")
        self.two_std_above = self.two_std_above * len(self.new_position_index_list)
        results_str = "data: %s," % (self.two_std_above)
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("data: [91, 91, 91, 91, 91, 91, 91],")
        self.two_std_below = self.two_std_below * len(self.new_position_index_list)
        results_str = "data: %s," % (self.two_std_below)
        my_html_test_doc[html_index] = results_str

        html_index = my_html_test_doc.index("<link href=\"bootstrap.min.css\" rel=\"stylesheet\">")
        bootstrap_string = "<link href=\""+Parameters.resources+"/css/bootstrap.min.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = bootstrap_string

        html_index = my_html_test_doc.index("<link href=\"dashboard.css\" rel=\"stylesheet\">")
        dashboard_string = "<link href=\""+ Parameters.resources+"/css/dashboard.css\" rel=\"stylesheet\">"
        my_html_test_doc[html_index] = dashboard_string

        f.close()
        html_file_name_margin = Parameters.results + \
                               "/html/html_leaf_margin_chart" + \
                               "_" + \
                               joined_final_l + \
                               ".html"
        new_html = open(html_file_name_margin, 'w')
        str_for_html = "".join(my_html_test_doc)
        new_html.write(str_for_html)
        # step four
        # output the html to file and open in a browser.
        filename = 'file:///' + html_file_name_margin

        # add success for this process to the config
        self.config_file.html_margin_output_config_file(html_file_name_margin)
        # Method##########

        # read in html
        method_html_test_doc = []
        method_file = open("../resources/html/method_index.html", "r")
        method_html_test_doc = method_file.read().split('\n')
        # substitute in method html
        method_index = method_html_test_doc.index("<title>Album example for Bootstrap</title>")
        method_str = "<title>Method</title>"
        method_html_test_doc[method_index] = method_str

        html_index = method_html_test_doc.index("<link href=\"bootstrap.min.css\" rel=\"stylesheet\">")
        bootstrap_string = "<link href=\"" + Parameters.resources + "/css/bootstrap.min.css\" rel=\"stylesheet\">"
        method_html_test_doc[html_index] = bootstrap_string

        html_index = method_html_test_doc.index("<link href=\"album.css\" rel=\"stylesheet\">")
        dashboard_string = "<link href=\"" + Parameters.resources + "/css/album.css\" rel=\"stylesheet\">"
        method_html_test_doc[html_index] = dashboard_string

        if "pdf.toJpg" in self.path:
            self.format = "PDF"
            self.format_conversion = "converted this to JPG"
        else:
            self.format = self.path[:-4]
            self.format_conversion = "not converted this"

        method_index = method_html_test_doc.index(
            "<p class=\"card-text\"><strong>Pre-processing - Step 1: Image File Extension Conversion.</strong><br />The original image file provided was in X format. MktStall has either converted/not converted this.</p>")
        method_str = "<p class=\"card-text\"><strong>Pre-processing - Step 1: Image File Extension Conversion.</strong><br />The original image file provided was in %s format. MktStall has %s.</p>" % (
            self.format, self.format_conversion)
        method_html_test_doc[method_index] = method_str

        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Image File Extension Conversion\">")
        method_str = "<img class=\"card-img-top\" src=\"%s\" alt=\"Image File Extension Conversion\">" % (
                self.path[:-4] + ".A.jpg")
        method_html_test_doc[method_index] = method_str

        # image orientation conversion
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Image Orientation Conversion\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Image Orientation Conversion\">" % self.path
        method_html_test_doc[method_index] = method_str

        method_index = method_html_test_doc.index(
            "<p class=\"card-text\"><strong>Pre-processing - Step 2: Image Orientation Conversion.</strong><br />The original image file was in X orientation. MktStall has either orientated this/not orientated this.</p>")
        method_str = "<p class=\"card-text\"><strong>Pre-processing - Step 2: Image Orientation Conversion.</strong><br />The original image file was in %s orientation. MktStall has %s this.</p>" % (
            self.orientation, self.rotated)
        method_html_test_doc[method_index] = method_str
        # segmentation
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Image Segmentation\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Image Segmentation\">" % (
                self.path + ".resizedforruler.jpg.masked.jpg")
        method_html_test_doc[method_index] = method_str

        # ruler:convert to binary
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Ruler: Convert to binary\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Ruler: Convert to binary\">" % (
                self.path + ".resizedforruler.B.H.jpg")
        method_html_test_doc[method_index] = method_str

        # ruler:contiguous pixels and bin
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Ruler: Count contiguous pixels and bin\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Ruler: Count contiguous pixels and bin\">" % (
                self.path + ".resizedforruler.jpg.masked.ruler_segmented.jpg")
        method_html_test_doc[method_index] = method_str

        # Ruler: From the top 2 bins with highest frequency, calculate mean pixels per block
        method_index = method_html_test_doc.index(
            "<p class=\"card-text\"><strong>Ruler - Step 3: Calculate mean pixels per block.</strong><br />From the top 2 bins with the highest frequency, MktStall calculates the mean pixels per block. We calculated this to be x pixels per centimeter.</p>")
        method_str = "<p class=\"card-text\"><strong>Ruler - Step 3: Calculate mean pixels per block.</strong><br />From the top 2 bins with the highest frequency, MktStall calculates the mean pixels per block. We calculated this to be %f pixels per centimeter.</p>" % pix_per_cm
        method_html_test_doc[method_index] = method_str

        # Label: Convert to grayscale
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Convert to grayscale\">")
        method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Convert to grayscale\">" % (
                self.path[:-4] + ".A.B.jpg")
        method_html_test_doc[method_index] = method_str

        # Label: Adaptative thresholding
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Adaptative thresholding\">")
        method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Adaptative thresholding\">" % (
                self.path[:-4] + ".A.B.C.jpg")
        method_html_test_doc[method_index] = method_str

        # Label: Crop
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Crop\">")
        method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Crop\">" % (
                self.path[:-4] + ".A.B.C.D.jpg")
        method_html_test_doc[method_index] = method_str

        # Label: Denoise with a Fast Non-local means algorithm

        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Denoise with a Fast Non-local means algorithm\">")
        method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Denoise with a Fast Non-local means algorithm\">" % (
                self.path[:-4] + ".A.B.C.D.E.jpg")
        method_html_test_doc[method_index] = method_str

        # Label: Progressively denoise - Global Thresholding
        # F = global_thresholding
        # G = otsus_thresholding
        # H = gaussian_filtering_otsus_thresholding
        # find if image is there
        global_threshold_file = self.path[:-4] + ".A.B.C.D.E.F.jpg"
        # add success for this process to the config
        last_working_denoiser_stage = "Fast Non-local means algorithm"
        if os.path.exists(global_threshold_file):
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise\">")
            method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Progressively denoise - Global thresholding\">" % (
                    self.path[:-4] + ".A.B.C.D.E.F.jpg")
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index("<p class=\"card-text\">Denoise5a</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5a: Progressively Denoise (Global Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The first stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Global Thresholding</a> the most lenient algorithm. MktStall performed this step on the image.</p>"
            method_html_test_doc[method_index] = method_str
            last_working_denoiser_stage = "Global Thresholding"
        else:
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise\">")
            method_str = "<img class=\"card-img-top\">"
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index(
                "<p class=\"card-text\">Denoise5a</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5a: Progressively Denoise (Global Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The first stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Global Thresholding</a> the most lenient algorithm. MktStall did not perform this step on the image as the previous step (%s) was sufficiently denoised.</p>" % last_working_denoiser_stage
            method_html_test_doc[method_index] = method_str

        # Label: Progressively denoise - Otsu's Thresholding
        otsus_thresholding_file = self.path[:-4] + ".A.B.C.D.E.G.jpg"
        if os.path.exists(otsus_thresholding_file):
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise - Otsu's Thresholding\">")
            method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Progressively denoise - Otsu's Thresholding\">" % (
                    self.path[:-4] + ".A.B.C.D.E.G.jpg")
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index("<p class=\"card-text\">Denoise5b</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5b: Progressively Denoise (Otsu's Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The second stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Otsu's Thresholding</a>, a slightly more stringent algorithm to the one used in step 5a. MktStall performed this step on the image.</p>"
            method_html_test_doc[method_index] = method_str

            last_working_denoiser_stage = "Otsu's Thresholding"
        else:
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise - Otsu's Thresholding\">")
            method_str = "<img class=\"card-img-top\">"
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index("<p class=\"card-text\">Denoise5b</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5b: Progressively Denoise (Otsu's Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The second stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Otsu's Thresholding</a>, a slightly more stringent algorithm to the one used in step 5a. MktStall did not perform this step on the image as the previous step (%s) was sufficiently denoised.</p>" % last_working_denoiser_stage
            method_html_test_doc[method_index] = method_str

        # Label: Progressively denoise - Gaussian Filtering and Otsu's Thresholding
        gaussian_otsus_thesholding_file = self.path[:-4] + ".A.B.C.D.E.H.jpg"
        if os.path.exists(gaussian_otsus_thesholding_file):
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise - Gaussian Filtering and Otsu's Thresholding\">")
            method_str = "<img class=\"card-img-top\" style=\"max-height:250px\" src=\"%s\" alt=\"Label: Progressively denoise - Gaussian Filtering and Otsu's Thresholding\">" % (
                    self.path[:-4] + ".A.B.C.D.E.H.jpg")
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index("<p class=\"card-text\">Denoise5c</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5c: Progressively Denoise (Gaussian Filtering and Otsu's Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The third stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Gaussian Filtering and Otsu's Thresholding</a>, the most stringent algorithm used. MktStall performed this step on the image.</p>"
            method_html_test_doc[method_index] = method_str

            last_working_denoiser_stage = "Gaussian Filtering and Otsu's Thresholding"
        else:
            method_index = method_html_test_doc.index(
                "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Label: Progressively denoise - Gaussian Filtering and Otsu's Thresholding\">")
            method_str = "<img class=\"card-img-top\">"
            method_html_test_doc[method_index] = method_str

            method_index = method_html_test_doc.index("<p class=\"card-text\">Denoise5c</p>")
            method_str = "<p class=\"card-text\"><strong>Label - Step 5c: Progressively Denoise (Gaussian Filtering and Otsu's Thresholding).</strong><br />MktStall progressively denoises the label. If the denoising algorithm is too lenient and does not sufficiently denoise the image, it can be too 'messy' to read, conversely, if the denoising algorithm is too stringent, the label cannot be easily read. The third stage of this progressive denoising step is <a href=\"https://docs.opencv.org/3.3.0/d7/d4d/tutorial_py_thresholding.html\">Gaussian Filtering and Otsu's Thresholding</a>, the most stringent algorithm used. MktStall did not perform this step on the image as the previous step (%s) was sufficiently denoised.</p>" % last_working_denoiser_stage
            method_html_test_doc[method_index] = method_str

        # Label - Apply OCR

        # Label - Label image
        method_index = method_html_test_doc.index(
            "<p class=\"card-text\"><strong>Label - Step 7: Label the image according to the label.</strong><br />After Optical Character recognition, MktStall has labelled the image according to the label:</p>")
        method_str = "<p class=\"card-text\"><strong>Label - Step 7: Label the image according to the label.</strong><br />After Optical Character recognition, MktStall has labelled the image according to the label: %s</p>" % final_l.encode(
            'utf-8')
        method_html_test_doc[method_index] = method_str

        # Stage: Denoise

        # Stage: Edge detection for all features

        # Stage: Select largest feature

        # Stage: Find centroid and contours

        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Stage: Find centroid and contours\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Stage: Find centroid and contours\">" % (
                path + "dilated_skinny.jpg")
        method_html_test_doc[method_index] = method_str
        # Stage: Crop area using centroid pm extremities of contours pm arbitary margin
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Stage: Crop area using centroid pm extremities of contours pm arbitary margin\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Stage: Crop area using centroid pm extremities of contours pm arbitary margin\">" % (
                self.path + ".cr.jpg")
        method_html_test_doc[method_index] = method_str

        # Stage: Reorientate leaf so major axis is vertical
        method_index = method_html_test_doc.index(
            "<img class=\"card-img-top\" data-src=\"holder.js/100px225?theme=thumb&bg=55595c&fg=eceeef&text=Thumbnail\" alt=\"Stage:Reorientate leaf so major axis is vertical\">")
        method_str = "<img class=\"card-img-top\" style=\"max-width:170px;height:auto;margin-left: auto;margin-right: auto;\" src=\"%s\" alt=\"Stage: Reorientate leaf so major axis is vertical\">" % (
                self.path + ".cr.jpg.ro.jpg")
        method_html_test_doc[method_index] = method_str

        # write to method html file
        html_file_name_method = Parameters.results + "/html/html_method" + "_" + joined_final_l + ".html"
        f_method = open(html_file_name_method, "w")
        method_html = "".join(method_html_test_doc)
        f_method.write(method_html)
        method_file_name = 'file:///' + html_file_name_method

        # add success for this process to the config
        self.config_file.html_method_output_config_file(html_file_name_method)
    @staticmethod
    def get_minimum_points_on_contour(distances, num_neighbours):
        """
        Builds a list of points that are a minimum in relations to the provided number of neighbours.
        :param distances: list of distances on the contour from the centroid.
        :param num_neighbours: the number of neighbours to test when looking for a minimum.
        :return: a list of points that are a minimum.
        """
        # Grab some memory to populate with minimum points.
        index_min = []
        # from number of neighbours to the length of the array
        # (less the number of neighbours so not to get index out of bounds error).
        for i in range(num_neighbours, (len(distances) - num_neighbours) - 1):
            count = 1;
            walking_left = True
            is_sinus_left = True
            while walking_left:
                if distances[i] > distances[i - count]:
                    is_sinus_left = False
                    walking_left = False
                if count == num_neighbours:
                    walking_left = False
                count = count + 1
            count = 1;
            walking_right = True;
            is_sinus_right = True
            while walking_right:
                if distances[i] > distances[i + count]:
                    is_sinus_right = False
                    walking_right = False
                if count == num_neighbours:
                    walking_right = False
                count = count + 1
            if is_sinus_left is True and is_sinus_right is True:
                index_min.append(i)
        return index_min

    @staticmethod
    def get_maximum_points_on_contour(distances, num_neighbours):
        """
        Builds a list of points that are a maximum in relation to the provided number of neigbours.
        :param distances: list of distances on the contour from the centroid.
        :param num_neighbours: the number of neighbours to test when looking for a maximum.
        :return: a list of points that are a maximum
        """
        index_max = []
        for i in range(num_neighbours, (len(distances) - num_neighbours) - 1):
            count = 1
            walking_left = True
            is_tip_left = True
            while walking_left:
                # do tip tests
                if distances[i] < distances[i - count]:
                    is_tip_left = False
                    walking_left = False
                if count == num_neighbours:
                    walking_left = False
                count = count + 1
            count = 1
            walking_right = True
            is_tip_right = True
            while walking_right:
                # do tip tests right
                if distances[i] < distances[i + count]:
                    is_tip_right = False
                    walking_right = False
                if count == num_neighbours:
                    walking_right = False
                count = count + 1
            if is_tip_left is True and is_tip_right is True:
                # append index max
                index_max.append(i)
        return index_max

    @staticmethod
    def remove_adjacent_local_maximum(distances, index_max, min_contour_adjacency):
        """
        Find local maximums that are adjacent and select the maximum that is furthest in distance from the centroid.
        :param distances: list of distances from the centroid on the contour.
        :param index_max: points that are local maximums
        :param min_contour_adjacency: the minimum number of positions allowed for to maximum points to exist
        :return: a refined list of local maximums
        """

        index_max_refined = []
        for i in range(1, len(index_max) - 2):
            if i == 0:
                bound_left = index_max[0] - min_contour_adjacency
                bound_right = index_max[0] + min_contour_adjacency
                left_position = (index_max[len(index_max) - 2]) - (len(distances) - 1)
                right_position = index_max[0 + 1]
                tip_most_position = index_max[0]
            elif i == len(index_max) - 2:
                bound_left = index_max[len(index_max) - 2] - min_contour_adjacency
                bound_right = index_max[len(index_max) - 2] + min_contour_adjacency
                left_position = index_max[len(index_max) - 2 - 1]
                right_position = (len(distances) - 1) + index_max[0]
                tip_most_position = index_max[len(index_max) - 2]
            else:
                bound_left = index_max[i] - min_contour_adjacency
                bound_right = index_max[i] + min_contour_adjacency
                left_position = index_max[i - 1]
                right_position = index_max[i + 1]
                tip_most_position = index_max[i]
            if left_position >= bound_left:
                if distances[left_position] > distances[tip_most_position]:
                    tip_most_position = left_position
            if right_position <= bound_right:
                if distances[right_position] > distances[tip_most_position]:
                    tip_most_position = right_position
            index_max_refined.append(tip_most_position)
        return index_max_refined

    def remove_neighbouring_local_minimums_overlapping_maximums(self, contour, index_minimums, index_maximums,
                                                                min_pos_dist):
        """
        Remove local minimums that are extra close to a maximum.
        :param index_maximums:
        :param index_minimums:
        :param contour:
        :param minimums: list of local minimums.
        :param maximums: list of local maximums.
        :param min_pos_dist: the minimum adjacency that a minimum can be to a maximum.
        :return:
        """

    def is_on_convex_hull(self, convex_hull, contour, contour_index, tolerance):
        """
        This function will test to see if the contour point indexed by contour_index sits
        on the convex hull perimeter within the tolerance.
        :param tolerance:
        :param convex_hull:
        :param contour:
        :param contour_index:
        :return:
        """
        # run through contour points
        # is it on convex hull

    @staticmethod
    def find_candidate_teeth(contour_tuples, minimums, maximums, centroid_image_path):
        """
        From a list of local minimums and a list of local maximums, find teeth triangles.
        :param centroid_image_path:
        :param contour_tuples:
        :param minimums: a list of local minimum.
        :param maximums: a list of local maximum.
        :return: a list of candidate teeth.
        """
        candidate_teeth = []
        position_tip = []
        position_sinus = []
        for i in range(0, len(minimums) - 1):
            for j in range(0, len(maximums) - 2):
                if maximums[j] < minimums[i] < maximums[j + 1]:
                    position_sinus.append(minimums[i])
        for i in range(0, len(maximums) - 1):
            for j in range(1, len(position_sinus) - 2):
                if position_sinus[j] < maximums[i] < position_sinus[j + 1]:
                    candidate_tooth = (position_sinus[j], maximums[i], position_sinus[j + 1])
                    position_tip.append(maximums[i])
                    candidate_teeth.append(candidate_tooth)
        # plot the candidate teeth onto the provided contour image.
        img = io.imread(centroid_image_path)
        for i in range(0, len(candidate_teeth) - 1):
            tooth_tuple = candidate_teeth[i]
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            cv2.circle(img, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img, sinus_a, tip, [0, 0, 255], 1)
            cv2.line(img, sinus_a, sinus_b, [0, 255, 0], 1)
            cv2.line(img, sinus_b, tip, [0, 0, 255], 1)
        io.imsave(centroid_image_path + ".t_cands.jpg", img)
        # add success for this process to the config
        return candidate_teeth

    def make_teeth_tuples(self, distance, positions_minimums, positions_maximums, coords):
        teeth = []
        for counter, value in enumerate(positions_maximums):
            for i in range(0, len(positions_minimums) - 1):
                if (positions_minimums[i] < value and value < positions_minimums[i + 1]):
                    tooth = [coords[positions_minimums[i]], coords[value], coords[positions_minimums[i + 1]]]
                    print tooth
                    teeth.append(tooth)
        return teeth

    def output_teeth_contour_tsv(self, centroid_img_path, contour_image_path, centroid, contour):
        # list of xy coordinate tuples on the contour.
        contour_tuples = []
        # list of distances from the centriod for each xy coordinate tuple on the contour.
        distances = []
        # list of indexes for contour tuples and distances.
        position_index_list = []
        ######################
        # Populate data structures.
        ######################
        orig_stdout = sys.stdout
        f = open(self.path + ".teeth_plot.tsv", 'w')
        sys.stdout = f
        # populate the position index list - this is simply the position on the contour for each xy coord tuple.
        for i in range(0, len(contour_tuples) - 1):
            position_index_list.append(i)
        # populate contour tuples as a list of tuples that are the contour.
        for tuple in contour:
            my_tuple = tuple[0]
            my_tuple_2 = (my_tuple[0], my_tuple[1])
            contour_tuples.append(my_tuple_2)
        # make a list of distances (from the centroid) for each of the tuples on the contour

        for tuple in contour_tuples:
            # distance = math.sqrt( math.pow(centroid[0]-tuple[0], 2) + math.pow(centroid[1] - tuple[1], 2) )
            distance = math.sqrt(math.pow(tuple[0] - centroid[0], 2) + math.pow(tuple[1] - centroid[1], 2))
            distances.append(distance)

        for i in range(0, len(contour_tuples) - 1):
            print i, '\t', contour_tuples[i], '\t', distances[i]

        sys.stdout = orig_stdout
        f.close()

    def points_in_poly(self, points, poly):
        polygon = Polygon(poly)
        for counter, value in enumerate(points):
            val = make_tuple(value)
            point = Point( float(int(val[0])), float(int(val[1]) ) )
            in_poly = polygon.contains(point)
            if in_poly:
                return True
        return False

    def tip_below_both_sinus(self, value):
        a = make_tuple(value[1])  # tip
        b = make_tuple(value[0])  # sinus_point_a
        c = make_tuple(value[2])  # sinus_point_b
        if a[1] >= b[1] and a[1] >= c[1]:
            return True
        else:
            return False

    def length_of_sinus_greater_than_width(self, value, centroid, leftmost, rightmost):
        b = make_tuple(value[0])  # sinus_point_a
        c = make_tuple(value[2])  # sinus_point_b
        length = self.get_line_length(b, c)

        length_left = self.get_line_length(centroid, leftmost)
        length_right = self.get_line_length(centroid, rightmost)
        half_width = (length_left + length_right) * 0.5
        if length > (half_width*0.5):
            return True
        else:
            return False


    def find_teeth_improved(self, centroid_img_path, contour_image_path, centroid, contour):

        # Make variables.
        number_of_teeth = 0
        position = []
        coords = []
        distance = []
        position_min = []
        position_max = []

        # load in the data from file.
        data_path = self.path + ".teeth_plot.tsv"
        with open(data_path, 'r') as csvfile:
            plots = csv.reader(csvfile, delimiter='\t')
            for row in plots:
                position.append(int(row[0]))
                coords.append(row[1])
                distance.append(float(row[2]))

        # calculate the mean and std deviation of the distance data.
        mean = np.mean(distance)
        stddev = np.std(distance)
        two_std_above = mean + (stddev * 2.0)
        two_std_below = mean - (stddev * 2.0)
        mean_x = (0,len(position)-1)
        mean_y = (mean, mean)
        two_std_below_x = (0,len(position)-1)
        two_std_below_y = (two_std_below,two_std_below)
        two_std_above_x = (0, len(position) - 1)
        two_std_above_y = (two_std_above, two_std_above)

        # Plot the distance, mean and std deviation lines.
        im = Image.open(centroid_img_path)
        draw = ImageDraw.Draw(im)
        plt.figure(num=None, figsize=(18, 18), dpi=80, facecolor='w', edgecolor='k')
        plt.subplot(211)
        plt.plot(position, distance)
        plt.xlabel("position on contour")
        plt.ylabel("distance from centroid")
        plt.title("Leaf contour plot")
        plt.plot(two_std_below_x[:], two_std_below_y[:], 'g-')
        plt.plot(two_std_above_x[:], two_std_above_y[:], 'g-')
        plt.plot(mean_x[:], mean_y[:], 'y-')

        # Generate a smoothend curve for  the distance plot.
        yhat = savgol_filter(distance, 51, 8, mode="nearest")  # window size 51, polynomial order 3
        plt.plot(position, yhat)

        # Using the smoothened curve, identify local maximums on the curve (candidate tooth tips)
        max_peakind = signal.argrelmax(np.asarray(yhat))
        for counter, value in enumerate(max_peakind[0]):
            if distance[value] < two_std_above and distance[value] > two_std_below:
                position_max.append(value)
                # Plot the local maxima (tooth tips)
                plt.plot(position[value], distance[value], 'ro')
            else:
                # Plot the local maxima above 2 std deviations (in green)
                plt.plot(position[value], distance[value], 'go')

        # Using the smoothened curve, identify local minimums on the curve (candidate tooth sinus)
        min_peakind = signal.argrelmin(np.asarray(yhat))
        for counter, value in enumerate(min_peakind[0]):
            if distance[value] < two_std_above and distance[value] > two_std_below:
                position_min.append(value)
                # Plot the local minima (tooth tips)
                plt.plot(position[value], distance[value], 'bo')
            else:
                # Plot the local minima below 2 std deviations (in green)
                plt.plot(position[value], distance[value], 'go')

        #########################
        # NOTE: signal.argrelmax produces a better results that signal.find_peaks_cwt.
        # Bioinformatics (2006) 22 (17): 2059-2065. doi: 10.1093/bioinformatics/btl355
        # http://bioinformatics.oxfordjournals.org/content/22/17/2059.long
        # max_peakind = signal.find_peaks_cwt(yhat, np.arange(1,30))
        # print max_peakind
        # for counter, value in enumerate(max_peakind):
        #    plt.plot(position[value], distance[value], 'ro')
        #    coord = make_tuple(coords[value])
        #    x = coord[0]
        #    y = coord[1]
        #    r = 1
        #    draw.ellipse((x-r, y-r, x+r, y+r), fill=(255,0,0,255))
        # generate an inverse numpy 1D arr (in order to find minima)
        # inv_yhat = []
        # for counter, value in enumerate(yhat):
        #    inv_yhat.append(-1.0*value)
        # minima : use builtin function fo find (min) peaks (use inversed data)
        # min_peakind = signal.find_peaks_cwt(inv_yhat, np.arange(1,30))
        # print min_peakind
        # for counter, value in enumerate(min_peakind):
        #    plt.plot(position[value], distance[value], 'bo')
        #    coord = make_tuple(coords[value])
        #    x = coord[0]
        #    y = coord[1]
        #    r = 1
        #    draw.ellipse((x-r, y-r, x+r, y+r), fill=(0,0,255,255))
        #########################

        # Make candidate teeth tuples that contains the triangle for each tooth.
        candidate_teeth = self.make_teeth_tuples(distance, position_min, position_max, coords)
        candidate_teeth_area = []
        pen_teeth_tuples = []
        post_teeth_tuples = []
        teeth_tuples = []
        filtered = []

        # Calculate the area for each of the candidate teeth.
        for counter, value in enumerate(candidate_teeth):
            a = make_tuple(value[1]) # tip
            b = make_tuple(value[0]) # sinus_point_a
            c = make_tuple(value[2]) # sinus_point_b
            x = 0
            y = 1
            area = 0.5 * abs((a[x] * b[y] + b[x] * c[y] + c[x] * a[y]) - (b[x] * a[y] + b[y] * c[x] + c[y] * a[x]))
            candidate_teeth_area.append(area)

        # Calculate the mean and std deviation for teeth area.
        area_mean = np.mean(candidate_teeth_area)
        area_stddev = np.std(candidate_teeth_area)
        teeth_area_two_std_above = area_mean + (area_stddev * 2.0)
        teeth_area_two_std_below = area_mean - (area_stddev * 2.0)

        # Generate flat and long bounding boxes for extreme topmost and bottomost points on contour.
        topmost = tuple(contour[contour[:, :, 1].argmin()][0])
        bottommost = tuple(contour[contour[:, :, 1].argmax()][0])
        leftmost = tuple(contour[contour[:, :, 0].argmin()][0])
        rightmost = tuple(contour[contour[:, :, 0].argmax()][0])
        height = 4
        height_bottom = 20 # half height of bounding box.
        width = 30
        topmost_bbox = [(leftmost[0], topmost[1] + height),(rightmost[0], topmost[1] + height),(rightmost[0], topmost[1] - height),(leftmost[0], topmost[1] - height)]
        bottommost_bbox = [(leftmost[0], bottommost[1] + height_bottom), (rightmost[0], bottommost[1] + height_bottom), (rightmost[0], bottommost[1] - height_bottom), (leftmost[0], bottommost[1] - height_bottom)]
        draw.polygon(topmost_bbox, outline=(0, 255, 0, 255))
        draw.polygon(bottommost_bbox, outline=(0, 255, 0, 255))


        # Filter out candidate teeth and place teeth that pass the filters into the final teeth list.
        for counter, value in enumerate(candidate_teeth):
            # Filter teeth that:
            #   have an area which is less than than the magic number from paper,
            #   have an area which is less than 2 std deviations from the mean teeth area,
            #   have an area which is greater than 2 std deviations from the mean teeth area,
            if candidate_teeth_area[counter] > self.magic_number_from_paper and \
                    candidate_teeth_area[counter] > teeth_area_two_std_below and \
                    self.points_in_poly(value, topmost_bbox) == False and \
                    self.points_in_poly(value, bottommost_bbox) == False and \
                    self.tip_below_both_sinus(value) == False and \
                    self.length_of_sinus_greater_than_width(value, centroid, leftmost, rightmost) == False:
                pen_teeth_tuples.append(value)
            else:
                filtered.append(value)

                #tip = make_tuple(value[1])
                #sinus_point_a = make_tuple(value[0])
                #sinus_point_b = make_tuple(value[2])
                #sinus = (sinus_point_a, sinus_point_b)
                #side_A = (tip, sinus_point_a)
                #side_B = (tip, sinus_point_b)
                #draw.line((sinus[0][0], sinus[0][1], sinus[1][0], sinus[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_A[0][0], side_A[0][1], side_A[1][0], side_A[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_B[0][0], side_B[0][1], side_B[1][0], side_B[1][1]), fill=(255, 255, 0, 255))
                #r = 1  # radius of an ellipse
                #x, y = sinus_point_a  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = sinus_point_b  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = tip  # red point which is a maximum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))


        # Filter out lobes.
        for counter, value in enumerate(pen_teeth_tuples):
            tip = make_tuple(value[1])
            sinus_a = make_tuple(value[0])
            sinus_b = make_tuple(value[2])
            # tip and centroid
            p1 = [tip[0], tip[1]]
            p2 = [centroid[0], centroid[1]]
            distance_tip_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            lobe_threshold = distance_tip_to_centroid * 0.25
            # sinus and centroid
            p1 = [sinus_a[0], sinus_a[1]]
            p2 = [centroid[0], centroid[1]]
            distance_sinus_a_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            difference_tcc_minus_sac = distance_tip_to_centroid - distance_sinus_a_to_centroid
            # sinusB and centroid
            p1 = [sinus_b[0], sinus_b[1]]
            p2 = [centroid[0], centroid[1]]
            distance_sinus_b_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            difference_tc_minus_sbc = distance_tip_to_centroid - distance_sinus_b_to_centroid
            if difference_tcc_minus_sac < lobe_threshold or difference_tc_minus_sbc < lobe_threshold:
                post_teeth_tuples.append(value)
            else:
                filtered.append(value)
                #tip = make_tuple(value[1])
                #sinus_point_a = make_tuple(value[0])
                #sinus_point_b = make_tuple(value[2])
                #sinus = (sinus_point_a, sinus_point_b)
                #side_A = (tip, sinus_point_a)
                #side_B = (tip, sinus_point_b)
                #draw.line((sinus[0][0], sinus[0][1], sinus[1][0], sinus[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_A[0][0], side_A[0][1], side_A[1][0], side_A[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_B[0][0], side_B[0][1], side_B[1][0], side_B[1][1]), fill=(255, 255, 0, 255))
                #r = 1  # radius of an ellipse
                #x, y = sinus_point_a  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = sinus_point_b  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = tip  # red point which is a maximum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                
        # Filter out small candidates based on triangle height
        for counter, value in enumerate(post_teeth_tuples):
            a = make_tuple(value[0])
            b = make_tuple(value[1])
            c = make_tuple(value[2])
            x = 0
            y = 1
            area = 0.5 * abs((a[x] * b[y] + b[x] * c[y] + c[x] * a[y]) - (b[x] * a[y] + b[y] * c[x] + c[y] * a[x]))
            sinus_a = a
            tip = b
            sinus_b = c
            lengths_between_sinus = self.get_line_length(sinus_a, sinus_b)
            height = 2 * (area / lengths_between_sinus)
            if height > self.magic_height_number:
                teeth_tuples.append(value)
            else:
                filtered.append(value)
                #tip = make_tuple(value[1])
                #sinus_point_a = make_tuple(value[0])
                #sinus_point_b = make_tuple(value[2])
                #sinus = (sinus_point_a, sinus_point_b)
                #side_A = (tip, sinus_point_a)
                #side_B = (tip, sinus_point_b)
                #draw.line((sinus[0][0], sinus[0][1], sinus[1][0], sinus[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_A[0][0], side_A[0][1], side_A[1][0], side_A[1][1]), fill=(255, 255, 0, 255))
                #draw.line((side_B[0][0], side_B[0][1], side_B[1][0], side_B[1][1]), fill=(255, 255, 0, 255))
                #r = 1  # radius of an ellipse
                #x, y = sinus_point_a  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = sinus_point_b  # blue point which is a minimum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))
                #x, y = tip  # red point which is a maximum
                #draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 255, 0, 255))

        # Draw accepted teeth outlines onto the final image for Leaf Teeth
        for counter, value in enumerate(teeth_tuples):
            tip = make_tuple(value[1])
            sinus_point_a = make_tuple(value[0])
            sinus_point_b = make_tuple(value[2])
            sinus = (sinus_point_a, sinus_point_b)
            side_A = (tip, sinus_point_a)
            side_B = (tip, sinus_point_b)
            draw.line((sinus[0][0], sinus[0][1], sinus[1][0], sinus[1][1]), fill=(0, 255, 0, 255))
            draw.line((side_A[0][0], side_A[0][1], side_A[1][0], side_A[1][1]), fill=(255, 0, 0, 255))
            draw.line((side_B[0][0], side_B[0][1], side_B[1][0], side_B[1][1]), fill=(255, 0, 0, 255))
            r = 1  # radius of an ellipse
            x, y = sinus_point_a  # blue point which is a minimum
            draw.ellipse((x - r, y - r, x + r, y + r), fill=(0, 0, 255, 255))
            x, y = sinus_point_b  # blue point which is a minimum
            draw.ellipse((x - r, y - r, x + r, y + r), fill=(0, 0, 255, 255))
            x, y = tip  # red point which is a maximum
            draw.ellipse((x - r, y - r, x + r, y + r), fill=(255, 0, 0, 255))
            # Count the number of teeth as we draw them.
            number_of_teeth += 1

        # Save the leaf teeth plot and image
        plt.savefig(centroid_img_path + '.teeth_plot.png', bbox_inches='tight')
        im.save(centroid_img_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_lobeF.jpg")

        return number_of_teeth

    def find_teeth(self, centroid_image_path, contour_image_path, centroid, contour):

        self.output_teeth_contour_tsv(centroid_image_path, contour_image_path, centroid, contour)
        return self.find_teeth_improved(centroid_image_path, contour_image_path, centroid, contour)

        # list of xy coordinate tuples on the contour.
        #contour_tuples = []
        # list of distances from the centriod for each xy coordinate tuple on the contour.
        #distances = []
        # list of indexes for contour tuples and distances.
        #position_index_list = []
        ######################
        # Populate data structures.
        ######################
        #orig_stdout = sys.stdout
        #f = open(self.path + ".teeth_plot.tsv", 'w')
        #sys.stdout = f
        # populate the position index list - this is simply the position on the contour for each xy coord tuple.
        #for i in range(0, len(contour_tuples) - 1):
        #    position_index_list.append(i)
        # populate contour tuples as a list of tuples that are the contour.
        #for tuple in contour:
        #    my_tuple = tuple[0]
        #    my_tuple_2 = (my_tuple[0], my_tuple[1])
        #    contour_tuples.append(my_tuple_2)
        # make a list of distances (from the centroid) for each of the tuples on the contour

        #for tuple in contour_tuples:
        #    # distance = math.sqrt( math.pow(centroid[0]-tuple[0], 2) + math.pow(centroid[1] - tuple[1], 2) )
        #    distance = math.sqrt(math.pow(tuple[0] - centroid[0], 2) + math.pow(tuple[1] - centroid[1], 2))
        #    distances.append(distance)

        #for i in range(0, len(contour_tuples) - 1):
        #    print i, '\t', contour_tuples[i], '\t', distances[i]

        #sys.stdout = orig_stdout
        #f.close()
        # testing
        #self.distances = distances
        #for i in range(0, len(distances) - 1):
        #    self.new_position_index_list.append(i)
        ######################
        # Find local minimum and maximum points on contour
        ######################
        # list of indexes into distances list for points that are a minimum on the contour.
        #index_min = self.get_minimum_points_on_contour(distances, 1)
        # list of indexes into distances list for points that are a maximum on the contour.
        #index_max = self.get_maximum_points_on_contour(distances, 2)
        # refine the local maximum by removing any adjacent/overlapping maximum
        #index_max_refined = self.remove_adjacent_local_maximum(distances, index_max, 5)
        ########################
        # Print convex hull to image
        ########################
        # for dev use only
        #hull = cv2.convexHull(contour)
        #index_max_refined_on_hull = []
        #img_convexhull = io.imread(centroid_image_path)
        #tolerance = 1
        #for index in index_max_refined:
        #    contour_point = contour_tuples[index]
        #    contour_point_x = contour_point[0]
        #    contour_point_y = contour_point[1]
        #    overlap = False
        #    for i in range(0, len(hull) - 1):
        #        hulltuple = hull[i][0]
        #        hulltuple_x = hulltuple[0]
        #        hulltuple_y = hulltuple[1]
        #        joined_hull_tuple = (hulltuple_x, hulltuple_y)
        #        if (hulltuple_x - tolerance) <= contour_point_x <= (hulltuple_x + tolerance) and \
        #                (hulltuple_y + tolerance) >= contour_point_y >= (hulltuple_y - tolerance):
        #            overlap = True
        #    if overlap:
        #        index_max_refined_on_hull.append(index)
        #    if not overlap:
        #        cv2.circle(img_convexhull, contour_point, 1, thickness=-1, color=(0, 0, 255))
        #for index in index_max_refined_on_hull:
        #    cv2.circle(img_convexhull, contour_tuples[index], 1, thickness=-1, color=(255, 0, 0))
        #min_pos_dist = 2
        #index_min_refined = []
        #for pos_min in index_min:
        #    overlap = False
        #    for pos_max in index_max_refined_on_hull:
        #        if (pos_max + min_pos_dist) > pos_min > (pos_max - min_pos_dist):
        #            overlap = True
        #    if not overlap:
        #        index_min_refined.append(pos_min)
        #io.imsave(centroid_image_path + ".cnvxhull.jpg", img_convexhull)
        # refine the local minimum by removing any adjacent/loverlamping with a maximum.
        # index_min_refined = self.remove_neighbouring_local_minimums_overlapping_maximums
        # (contour, index_min, index_max_refined, 1)
        # plot the minimums and maximums of a distance plot.
        #PeaUtils.show_leaf_teeth_distance_plot("Index Refined Tips and Sinus", "position", "distance", distances,
        #                                            index_max_refined, "index_max - tips", index_min_refined,
        #                                            "index_min - sinus")
        # plot the minimums and maximums on a centroid image of the leaf.
        #PeaUtils.output_image_points_for_leaf_teeth(centroid_image_path, contour_tuples, index_max_refined,
        #                                                index_min_refined, "refined_points")
        ########################
        # Find candidate teeth.
        ########################
        # from the refined minimums and maximums find candidate teeth.
        # candidate_teeth = self.find_candidate_teeth(contour_tuples, index_min_refined, index_max_refined,
        #                                            centroid_image_path)
        # number_of_teeth = self.filter_candidate_teeth(candidate_teeth, contour_tuples, centroid_image_path, distances,
        #                                              contour, centroid)
        # return number_of_teeth

    def filter_candidate_teeth(self, candidate_teeth, contour_tuples, centroid_image_path, distances, contour,
                               centroid):
        ########################
        # Filter teeth based on area.
        ########################
        # Calculate teeth area and filter teeth based on area
        candidate_teeth_area = []
        for i in candidate_teeth:
            a = contour_tuples[i[0]]
            b = contour_tuples[i[1]]
            c = contour_tuples[i[2]]
            x = 0
            y = 1
            area = 0.5 * abs((a[x] * b[y] + b[x] * c[y] + c[x] * a[y]) - (b[x] * a[y] + b[y] * c[x] + c[y] * a[x]))
            candidate_teeth_area.append(area)
        area_filtered_candidate_teeth = []
        img6 = io.imread(centroid_image_path)
        for i in range(0, len(candidate_teeth) - 1):
            if candidate_teeth_area[i] > self.magic_number_from_paper:
                tooth_tuple = candidate_teeth[i]
                area_filtered_candidate_teeth.append(candidate_teeth[i])
                sinus_a = contour_tuples[tooth_tuple[0]]
                tip = contour_tuples[tooth_tuple[1]]
                sinus_b = contour_tuples[tooth_tuple[2]]
                cv2.circle(img6, sinus_a, 1, thickness=-1, color=(0, 0, 255))
                cv2.circle(img6, tip, 1, thickness=-1, color=(0, 255, 0))
                cv2.circle(img6, sinus_b, 1, thickness=-1, color=(0, 0, 255))
                cv2.line(img6, sinus_a, tip, [0, 0, 255], 1)
                cv2.line(img6, sinus_a, sinus_b, [0, 255, 0], 1)
                cv2.line(img6, sinus_b, tip, [0, 0, 255], 1)
        #plt.savefig(str(Parameters.results) + '/images/teeth_contour_plot.png')
        io.imsave(centroid_image_path + ".area_teeth.jpg", img6)
        # add success for this process to the config
        self.config_file.area_teeth_output_config_file(centroid_image_path + ".area_teeth.jpg")
        #######################
        # Filter teeth that have a tip more than x std-devs from the mean.
        #######################
        # calculate mean and std deviation
        mean = np.mean(distances)
        # testing
        self.mean_dist_from_centroid.append(mean)
        #
        stddev = np.std(distances)
        multiplyer_std_above = mean + (self.std_multiplier * stddev)
        multiplyer_std_below = mean - (self.std_multiplier * stddev)
        # testing
        one_std_above = mean + (stddev)
        one_std_below = mean - (stddev)
        two_std_above = mean + (stddev * 2)
        two_std_below = mean - (stddev * 2)
        self.onestd_multiplierabove.append(multiplyer_std_above)
        self.onestd_multiplierbelow.append(multiplyer_std_below)
        self.one_std_above.append(one_std_above)
        self.one_std_below.append(one_std_below)
        self.two_std_above.append(two_std_above)
        self.two_std_below.append(two_std_below)
        #
        # teeth above 2sd above the mean
        # new list area_filtered std_deviation_filtered_cand_teeth
        area_std_dev_filt_cand_teeth = []
        # loop area filteredcandidate teeth
        for tuple in area_filtered_candidate_teeth:
            if multiplyer_std_above > distances[tuple[1]] > multiplyer_std_below:
                area_std_dev_filt_cand_teeth.append(tuple)
        img7 = io.imread(centroid_image_path)
        # for i in range(0, len(area_std_dev_filt_cand_teeth) - 1):
        for tooth_tuple in area_std_dev_filt_cand_teeth:
            # tooth_tuple = area_std_dev_filt_cand_teeth[i]
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            cv2.circle(img7, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img7, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img7, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img7, sinus_a, tip, [0, 0, 255], 1)
            cv2.line(img7, sinus_a, sinus_b, [0, 255, 0], 1)
            cv2.line(img7, sinus_b, tip, [0, 0, 255], 1)
        io.imsave(centroid_image_path + ".area_teeth_2sd.jpg", img7)
        # add success for this process to the config
        self.config_file.area_teeth_2sd_output_config_file(centroid_image_path + ".area_teeth_2sd.jpg")
        # Filter teeth based on the length of the line between sinus a and sinus b of the tooth.
        ##############################
        area_std_dev_filt_cand_teeth = PeaUtils.get_unique(area_std_dev_filt_cand_teeth)
        lengths_between_sinus = []
        sum_length_for_adjacent_sides = []
        for tooth_tuple in area_std_dev_filt_cand_teeth:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            lengths_between_sinus.append(self.get_line_length(sinus_a, sinus_b))
            sum_length_for_adjacent_sides.append(self.get_line_length(sinus_a, tip) + self.get_line_length(tip, sinus_b))

        area_std_dev_filt_cand_teeth_sinus_length = []
        index = 0
        list_of_differences_opposite_side_v_sum_of_adjacent = []
        for tooth_tuple in area_std_dev_filt_cand_teeth:
            difference = sum_length_for_adjacent_sides[index] - lengths_between_sinus[index]
            list_of_differences_opposite_side_v_sum_of_adjacent.append(difference)
            if (sum_length_for_adjacent_sides[index] - lengths_between_sinus[
                index]) > self.a_magic_number_for_inequality:
                area_std_dev_filt_cand_teeth_sinus_length.append(tooth_tuple)
            index = index + 1
        mean_for_list_of_differences_opposite_side_v_sum_of_adjacent = np.mean(
            list_of_differences_opposite_side_v_sum_of_adjacent)
        std_dev_for_mean_for_list_of_differences_opposite_side_v_sum_of_adjacent = np.std(
            list_of_differences_opposite_side_v_sum_of_adjacent)
        img8 = io.imread(centroid_image_path)
        for tooth_tuple in area_std_dev_filt_cand_teeth_sinus_length:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            cv2.circle(img8, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img8, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img8, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img8, sinus_a, tip, [0, 0, 255], 2)
            cv2.line(img8, sinus_a, sinus_b, [0, 255, 0], 2)
            cv2.line(img8, sinus_b, tip, [0, 0, 255], 2)
        io.imsave(centroid_image_path + ".area_teeth_2sd_sinusPointsLength.jpg", img8)
        # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_output_config_file(centroid_image_path + ".area_teeth_2sd_sinusPointsLength.jpg")
        ##############################
        # Filter teeth that have a hight that is too high to be a tooth.
        ##############################
        area_std_dev_filt_cand_teeth_sinus_length_height = []
        index = 0
        for tooth_tuple in area_std_dev_filt_cand_teeth_sinus_length:
            a = contour_tuples[tooth_tuple[0]]
            b = contour_tuples[tooth_tuple[1]]
            c = contour_tuples[tooth_tuple[2]]
            x = 0
            y = 1
            area = 0.5 * abs((a[x] * b[y] + b[x] * c[y] + c[x] * a[y]) - (b[x] * a[y] + b[y] * c[x] + c[y] * a[x]))
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            lengths_between_sinus = self.get_line_length(sinus_a, sinus_b)
            height = 2 * (area / lengths_between_sinus)
            if height > self.magic_height_number:
                area_std_dev_filt_cand_teeth_sinus_length_height.append(tooth_tuple)
            index = index + 1
        img9 = io.imread(centroid_image_path)
        area_std_dev_filt_cand_teeth_sinus_length_height = PeaUtils.get_unique(
            area_std_dev_filt_cand_teeth_sinus_length_height)

        triangle_geom_list = []
        triangle_tuple_list = []
        for tooth_tuple in area_std_dev_filt_cand_teeth_sinus_length_height:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            cv2.circle(img9, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img9, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img9, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img9, sinus_a, tip, [0, 0, 255], 1)
            cv2.line(img9, sinus_a, sinus_b, [0, 255, 0], 1)
            cv2.line(img9, sinus_b, tip, [0, 0, 255], 1)
            tooth_triangle = Polygon([sinus_a, tip, sinus_b])
            triangle_geom_list.append(tooth_triangle)
            triangle_tuple_list.append(tooth_tuple)
        io.imsave(centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height.jpg", img9)
        # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_height_output_config_file(centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height.jpg")
        self.number_of_teeth = len(area_std_dev_filt_cand_teeth_sinus_length_height)

        key = 0
        triangle_dictionary = {}
        tuple_dictionary = {}
        for index in range(0, len(triangle_geom_list)):
            triangle = triangle_geom_list[index]
            tuple = triangle_tuple_list[index]
            triangle_dictionary.update({key: triangle})
            tuple_dictionary.update({key: tuple})
            key = key + 1

        key_a = 0
        for index_a in range(0, len(triangle_geom_list)):
            img10 = io.imread(centroid_image_path)
            sinus_a = contour_tuples[triangle_tuple_list[index_a][0]]
            tip = contour_tuples[triangle_tuple_list[index_a][1]]
            sinus_b = contour_tuples[triangle_tuple_list[index_a][2]]
            cv2.circle(img10, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img10, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img10, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img10, sinus_a, tip, [0, 0, 255], 1)
            cv2.line(img10, sinus_a, sinus_b, [0, 255, 0], 1)
            cv2.line(img10, sinus_b, tip, [0, 0, 255], 1)
            temp_out_path = ""
            temp_out_path = centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_ID_key_a.jpg"
            io.imsave(temp_out_path, img10)
            
            # print "#######OUTER-LOOP#######"
            # print "keyA",keyA
            key_b = key_a + 1
            if key_b <= len(triangle_geom_list):

                for indexB in range(key_b, len(triangle_geom_list)):
                    # print "###inner-loop"
                    # print "keyB", keyB
                    triangle_a = triangle_geom_list[index_a]
                    triangle_b = triangle_geom_list[indexB]
                    if key_b in triangle_dictionary and key_a != key_b and triangle_a.intersects(triangle_b):
                        triangle_a_area = triangle_a.area
                        triangle_b_area = triangle_b.area
                        intersection_triangles = triangle_a.intersection(triangle_b)
                        intersection_area = intersection_triangles.area
                        if triangle_a_area < triangle_b_area:
                            smallest = triangle_a
                            # print "Smallest Triangle Key is A: "
                            # print "Smallest Triangle Key: ",keyA
                            # print "Smallest Triangle: ",triangleA
                            # print "Smallest Triangle Area: ",triangleAarea
                            # print "Largest Triangle: ",triangleB
                            # print "Largest Triangle Area: ",triangleBarea
                        else:
                            smallest = triangle_b
                            # print "Smallest Triangle Key is B: "
                            # print "Smallest Triangle Key: ",keyB
                            # print "Smallest Triangle: ",triangleB
                            # print "Smallest Triangle Area: ",triangleBarea
                            # print "Largest Triangle: ",triangleA
                            # print "Largest Triangle Area: ",triangleAarea

                        # print "Intersection Area: ",intersectionarea
                        max_overlap = smallest.area * 0.3

                        # print "Maximum allowable overlap (smallest area * 0.03): ",max_overlap
                        if intersection_area > (smallest.area * 0.3):
                            # print "Maximum allowable overlap exceeded!"
                            # print "Removing ",keyB," from dictionary."
                            del triangle_dictionary[key_b]
                            del tuple_dictionary[key_b]
                    key_b = key_b + 1
            key_a = key_a + 1


        triangle_overlap_filtered = []
        for key in tuple_dictionary:
            triangle_overlap_filtered.append(tuple_dictionary[key])

        img10 = io.imread(centroid_image_path)
        for tooth_tuple in triangle_overlap_filtered:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            cv2.circle(img10, sinus_a, 1, thickness=-1, color=(0, 0, 255))
            cv2.circle(img10, tip, 1, thickness=-1, color=(0, 255, 0))
            cv2.circle(img10, sinus_b, 1, thickness=-1, color=(0, 0, 255))
            cv2.line(img10, sinus_a, tip, [0, 0, 255], 1)
            cv2.line(img10, sinus_a, sinus_b, [0, 255, 0], 1)
            cv2.line(img10, sinus_b, tip, [0, 0, 255], 1)
        io.imsave(centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered.jpg", img10)


        # find overlapping teeth, keep one tooth

        apex_filtered = []
        # REMOVE APEX COUNTED AS TEETH
        apex_y = self.get_apex_y_coord(contour)
        img10 = io.imread(centroid_image_path)
        for tooth_tuple in triangle_overlap_filtered:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]
            if tip[1] > apex_y + 4:
                apex_filtered.append(tooth_tuple)
                cv2.circle(img10, sinus_a, 1, thickness=-1, color=(0, 0, 255))
                cv2.circle(img10, tip, 1, thickness=-1, color=(0, 255, 0))
                cv2.circle(img10, sinus_b, 1, thickness=-1, color=(0, 0, 255))
                cv2.line(img10, sinus_a, tip, [0, 0, 255], 1)
                cv2.line(img10, sinus_a, sinus_b, [0, 255, 0], 1)
                cv2.line(img10, sinus_b, tip, [0, 0, 255], 1)
                temp_out_path = centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_apexF.jpg"
                #TODO: the above is not made
                ###############
                io.imsave(temp_out_path, img10)
        # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_apexF_output_config_file(
            centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_apexF.jpg")
        # lobe detection
        lobe_filtered = []
        # REMOVE lobe COUNTED AS TEETH
        img10 = io.imread(centroid_image_path)
        for tooth_tuple in apex_filtered:
            sinus_a = contour_tuples[tooth_tuple[0]]
            tip = contour_tuples[tooth_tuple[1]]
            sinus_b = contour_tuples[tooth_tuple[2]]

            # tip and centroid
            p1 = [tip[0], tip[1]]
            p2 = [centroid[0], centroid[1]]
            distance_tip_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            lobe_threshold = distance_tip_to_centroid * 0.25

            # sinus and centroid
            p1 = [sinus_a[0], sinus_a[1]]
            p2 = [centroid[0], centroid[1]]
            distance_sinus_a_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            difference_tcc_minus_sac = distance_tip_to_centroid - distance_sinus_a_to_centroid

            # sinusB and centroid
            p1 = [sinus_b[0], sinus_b[1]]
            p2 = [centroid[0], centroid[1]]
            distance_sinus_b_to_centroid = math.sqrt(((p1[0] - p2[0]) ** 2) + ((p1[1] - p2[1]) ** 2))
            difference_tc_minus_sbc = distance_tip_to_centroid - distance_sinus_b_to_centroid
            if difference_tcc_minus_sac < lobe_threshold or difference_tc_minus_sbc < lobe_threshold:
                lobe_filtered.append(tooth_tuple)
                cv2.circle(img10, sinus_a, 1, thickness=-1, color=(0, 0, 255))
                cv2.circle(img10, tip, 1, thickness=-1, color=(0, 255, 0))
                cv2.circle(img10, sinus_b, 1, thickness=-1, color=(0, 0, 255))
                cv2.line(img10, sinus_a, tip, [0, 0, 255], 1)
                cv2.line(img10, sinus_a, sinus_b, [0, 255, 0], 1)
                cv2.line(img10, sinus_b, tip, [0, 0, 255], 1)
                temp_out_path = ""
                temp_out_path = centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_lobeF.jpg"
                # TODO: the above is not made
                #######
                io.imsave(temp_out_path, img10)
                # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_lobeF_output_config_file(
            centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_lobeF.jpg")
        print "Number of teeth: ", len(area_std_dev_filt_cand_teeth_sinus_length_height)
        print "New Number of teeth: ", len(lobe_filtered)
        # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_ID_key_a_output_config_file(
            centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_ID_key_a.jpg")
        # add success for this process to the config
        self.config_file.area_teeth_2sd_sinusPointsLength_height_overlapsFiltered_output_config_file(
            centroid_image_path + ".area_teeth_2sd_sinusPointsLength_height_overlapsFiltered.jpg")
        self.number_of_teeth = len(triangle_overlap_filtered)
        return len(lobe_filtered)

    @staticmethod
    def get_apex_y_coord(contour):
        x, y, w, h = cv2.boundingRect(contour)
        return y

    @staticmethod
    def point_within_bounding_box(min_row, min_col, max_row, max_col):
        # if min_row > 100 and min_col > 100 and max_row < 600 and max_col < 380:
        if min_row > 200 and min_col > 200 and max_row < 850 and max_col < 1200:
            return "true"

    @staticmethod
    def measure_properties(region):
        print region.minor_axis_length / region.major_axis_length, region.area, region.eccentricity, region.convex_area
        # print diameter
        print("major axis %f" % region.minor_axis_length)
        print("minor axis %f" % region.major_axis_length)
        print("area %f" % region.area)
        print("perimeter %f" % region.perimeter)
        # print("centroid" + region.local_centroid)
        print("aspect ratio(major:minor) %f" % (region.major_axis_length / region.minor_axis_length))

        print("perimeter ratio of major and minor axis length %f" % (
                region.perimeter / (region.minor_axis_length + region.major_axis_length)))
        # print("convex area %f"% region.convex_area)
        # print("perimeter convexity %f" % region.convex_area.perimeter)
        print("area convexity %f" % ((region.convex_area - region.area) / region.area))
        print("solidity %f" % (region.area / region.convex_area))
        # sphericity
        # print("equivalent diameter %s" (region.equivalent_diameter))

        print("roundness %f" % ((4 * math.pi * region.area) / (region.perimeter * region.perimeter)))
        print("compactness %f" % ((region.perimeter * region.perimeter) / region.area))
        print("rectangularity %f" % (region.area / (region.minor_axis_length * region.major_axis_length)))
        print("eccentricity %f" % region.eccentricity)
        # print("narrow factor")
        # print("perimeter ratiof of diamter")
        # print("perimeter ratio of major axis %f"% (region.perimeter/region.major_axis_length))

        # print("ellipse variance")
        # print("smooth factor")
        # leaf width factor
        # area width factor

    def extract_image_ruler(self):
        print "extract_image_ruler: ", self.path
        reference_image = io.imread(self.path)
        resized_image_ruler = rescale(reference_image, 0.5, mode='constant')  # 50% more pixels
        io.imsave(self.path + ".resizedforruler.jpg", resized_image_ruler)
        self.image_ruler = RulerMeasure(self.path + ".resizedforruler.jpg", self.config_file)
        # add success for this process to the config
        self.config_file.resizedforruler_output_config_file(self.path + ".resizedforruler.jpg")
        self.image_ruler.process_ruler()

    def extract_image_label(self):
        """
        Function: extracts the image label from this image and assigns it to this image.
        """
        if self.DEBUG:
            print ""
            print "#################################################"
            print "Starting extract_image_label for: ", self.name
            print "Performing first stage extraction upto denoised image (E)."

        ############################
        rotated_image = self.rotate_image_if_portrait_copy(self.name, self.path)
        # add success for this process to the config - A
        self.config_file.A_output_config_file(rotated_image.path)
        greyscale_image_object = self.get_greyscale_copy(rotated_image.name, rotated_image.path, self.config_file)
        # add success for this process to the config - B
        self.config_file.B_output_config_file(greyscale_image_object.path)
        adaptive_threshold_image_object = self.get_adpative_threshold_copy(greyscale_image_object.name,
                                                                           greyscale_image_object.path, self.config_file)
        # add success for this process to the config - C
        self.config_file.C_output_config_file(adaptive_threshold_image_object.path)
        crop_image_object = self.get_cropped_copy(adaptive_threshold_image_object.name,
                                                  adaptive_threshold_image_object.path, 30, 40, self.config_file)
        # add success for this process to the config - D
        self.config_file.D_output_config_file(crop_image_object.path)
        denoised_image_object = self.get_denoised_copy(crop_image_object.name, crop_image_object.path, self.config_file)
        # add success for this process to the config - E
        self.config_file.E_output_config_file(denoised_image_object.path)
        img = cv2.imread(denoised_image_object.path, 0)
        print "Getting edges."
        edges = cv2.Canny(img, 100, 200)
        #print "Plotting Original vs Edge images."
        #plt.subplot(121), plt.imshow(img, cmap='gray')
        #print "Making original image."
        #plt.title('Original Image'), plt.xticks([]), plt.yticks([])
        #plt.subplot(122), plt.imshow(edges, cmap='gray')
        #print "Making edge image."
        #plt.title('Edge Image'), plt.xticks([]), plt.yticks([])
        #plt.savefig(Parameters.results + '/images/testingedge.png')
        #plt.show()
        print "Getting kernel and dilated labels."
        kernel_label = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        dilated_label = cv2.dilate(edges, kernel_label)
        print "Show img with dilated edges on canny edge detection."
        new_contours, hierarchy = cv2.findContours(dilated_label.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
        contours_found = sorted(new_contours, key=cv2.contourArea, reverse=True)[:10]
        contour = contours_found[0]

        # largest_areas = sorted(contour, key=cv2.contourArea)
        # cv2.drawContours(label_for_drawing_1, [largest_areas[-2]], -1, (0, 0, 255), 2)
        label_for_drawing_1 = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)

        # M = cv2.moments(contour)
        # leftmost = tuple(contour[contour[:, :, 0].argmin()][0])
        # rightmost = tuple(contour[contour[:, :, 0].argmax()][0])

        cv2.drawContours(label_for_drawing_1, contour, -1, (0, 0, 255), 2)
        # PeaUtils.drawline(label_for_drawing_1, leftmost, rightmost,  (0, 0, 255), 2)
        cv2.imwrite(self.path + ".labelcontours.jpg", label_for_drawing_1)

        # rect = cv2.minAreaRect(contour)
        # box = np.int0(cv2.cv.BoxPoints(rect))
        # cv2.drawContours(label_for_drawing_1, [box], 0, (0, 0, 255), 2)

        # cv2.imwrite(self.path + ".labelcontours.jpg", label_for_drawing_1)
        # (x, y), (MA, ma), angle = cv2.fitEllipse(contour)
        #
        # if angle > 90 and angle < 270:
        #     if angle > 90 and angle <= 180:
        #         angle = -(180.0 - angle)
        #     if angle > 180 and angle <= 270:
        #         angle = angle - 180
        # rows, cols, channels = image.shape
        # M = cv2.getRotationMatrix2D((x, y), angle, 1)
        # rotated_img = cv2.warpAffine(image, M, (cols, rows))
        # saving
        # draw contouurs
        # label_for_drawing_1 = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
        # cv2.drawContours(label_for_drawing_1, contour, -1, (0, 0, 255), 2)
        # cv2.imwrite(self.path + ".labelcontours.jpg", label_for_drawing_1)
        ############################
        # KEEP ME
        # rotated_image = self.rotate_image_if_portrait_copy(self.name, self.path)
        # greyscale_image_object = self.get_greyscale_copy(rotated_image.name, rotated_image.path)
        # adaptivethreshold_image_object = self.get_adpative_threshold_copy(greyscale_image_object.name,
        #                                                                  greyscale_image_object.path)
        # crop_image_object = self.get_cropped_copy(adaptivethreshold_image_object.name,
        #                                           adaptivethreshold_image_object.path, 30, 40)
        # denoised_image_object = self.get_denoised_copy(crop_image_object.name, crop_image_object.path)
        #########
        # orientate conto so it is horizontal
        tmp_image_label = imagelabel.ImageLabel(
            self.extract_text_ocr(denoised_image_object.name, denoised_image_object.path))
        accepted = tmp_image_label.process_image_label()
        if accepted:
            self.image_label = tmp_image_label
            if self.DEBUG:
                print "Image label accepted: ", accepted
                print self.image_label.print_image_label_list("ImageLabelList: E: ")
                print "Success at E"
                print "#################################################"
                print ""
            if self.SHOW_IMAGES:
                PeaUtils.show_two_images("crop_image_object", "denoised_image_object", crop_image_object,
                                                  denoised_image_object)
        else:
            if self.DEBUG:
                print "Image label accepted: ", accepted
                print "Performing second stage extraction (F)."
            global_threshold_image_object = self.global_thresholding(denoised_image_object.name,
                                                                    denoised_image_object.path, self.config_file)
            # add success for this process to the config - F
            self.config_file.F_output_config_file(global_threshold_image_object.path)
            tmp_image_label = imagelabel.ImageLabel(
                self.extract_text_ocr(global_threshold_image_object.name, global_threshold_image_object.path))
            accepted = tmp_image_label.process_image_label()
            if accepted:
                self.image_label = tmp_image_label
                if self.DEBUG:
                    print "Image label accepted: ", accepted
                    print self.image_label.print_image_label_list("ImageLabelList: F: ")
                    print "Success at F"
                    print "#################################################"
                    print ""
                if self.SHOW_IMAGES:
                    PeaUtils.show_two_images("denoised_image_object", "globalthreshold_image_object",
                                                      denoised_image_object, global_threshold_image_object)
            else:
                if self.DEBUG:
                    print "Image label accepted: ", accepted
                    print "Performing third stage extraction (G)."
                otsus_thresholding_image_object = self.otsus_thresholding(denoised_image_object.name,
                                                                          denoised_image_object.path, self.config_file)
                # ADD SUCCESS HERE at G
                # add success for this process to the config - G
                self.config_file.G_output_config_file(otsus_thresholding_image_object.path)
                tmp_image_label = imagelabel.ImageLabel(
                    self.extract_text_ocr(otsus_thresholding_image_object.name, otsus_thresholding_image_object.path))
                accepted = tmp_image_label.process_image_label()
                if accepted:
                    self.image_label = tmp_image_label
                    if self.DEBUG:
                        print "Image label accepted: ", accepted
                        print self.image_label.print_image_label_list("ImageLabelList: G: ")
                        print "Success at G"
                        print "#################################################"
                        print ""
                    if self.SHOW_IMAGES:
                        PeaUtils.show_two_images("denoised_image_object", "otsus_thresholding_image_object",
                                                          denoised_image_object, otsus_thresholding_image_object)
                else:
                    if self.DEBUG:
                        print "Image label accepted: ", accepted
                        print "Performing fourth stage extraction (H)."
                    gaussian_filtering_otsus_thresholding_image_object = self.gaussian_filtering_otsus_thresholding(
                        denoised_image_object.name, denoised_image_object.path, self.config_file)
                    # Add succss at H
                    # add success for this process to the config - G
                    self.config_file.H_output_config_file(gaussian_filtering_otsus_thresholding_image_object.path)
                    tmp_image_label = imagelabel.ImageLabel(
                        self.extract_text_ocr(gaussian_filtering_otsus_thresholding_image_object.name,
                                              gaussian_filtering_otsus_thresholding_image_object.path))

                    accepted = tmp_image_label.process_image_label()
                    if accepted:
                        self.image_label = tmp_image_label
                        if self.DEBUG:
                            print "Image label accepted: ", accepted
                            print self.image_label.print_image_label_list("ImageLabelList: H: ")
                            print "Success at H"
                            print "#################################################"
                            print ""
                        if self.SHOW_IMAGES:
                            PeaUtils.show_two_images("denoised_image_object",
                                                              "gaussian_filtering_otsus_thresholding_image_object",
                                                              denoised_image_object,
                                                              gaussian_filtering_otsus_thresholding_image_object)
                    else:

                        tmp_image_label.imageLabelList = [PeaImage.get_clean_ID(self.name), ".no_label"]
                        self.image_label = tmp_image_label
                        if self.DEBUG:
                            print "Image label accepted: ", accepted
                            print "No success - unable to read label"
                            print "#################################################"
                            print ""
                        if self.SHOW_IMAGES:
                            PeaUtils.show_two_images("denoised_image_object",
                                                              "gaussian_filtering_otsus_thresholding_image_object",
                                                              denoised_image_object,
                                                              gaussian_filtering_otsus_thresholding_image_object)
        print "BEFORE YOU OUTPUT STRING"
        print "Image: ", self.name
        print "Label: ", self.image_label.imageLabelList
        print "Label Accepted: ", self.image_label.imageAcceptable

        #self.config_file.label_string_output_config_file(self.image_label)
        if self.VERBOSE:
            print "Image: ", self.name
            print "Label: ", self.image_label.imageLabelList
            print "Label Accepted: ", self.image_label.imageAcceptable

    @staticmethod
    def get_clean_ID(in_name):
        cleanOne = in_name.replace("pdf.toJpg.","")
        cleanTwo = cleanOne.replace(".jpg","")
        cleanThree = cleanTwo.replace("_.",".")
        return cleanTwo

    @staticmethod
    def extract_text_ocr(in_name, in_path):
        """
        Function: Performs Optical Character Recognition to read the label
        :param in_name: Name of image
        :param in_path: Path of image
        :return:
        """
        image_containing_text = Image.open(in_path)
        image_containing_text.load()
        image_containing_text.split()
        text_list = pytesseract.image_to_string(image_containing_text).splitlines()  # read the label and place each line in a list
        return text_list

    def rotate_image_if_portrait_copy(self, in_name, in_path):
        """
        Function rotates the provided disk image if needed, otherwise the disk_image provided is returned.
        :return: rotated disk_image or just disk_image if already landscape.
        """
        img = Image.open(in_path)
        dimensions = img.size
        img.load()
        outname, outpath = PeaUtils.get_file_path("A", in_path, ".jpg")
        if dimensions[0] < dimensions[1]:
            # width 0 height 1
            # if width < height it is portrait
            self.orientation = "Portrait"
            img_landscape = img.transpose(PIL.Image.ROTATE_90)
            img_landscape.save(outpath)
            self.rotated = "rotated"
        else:
            self.orientation = "Landscape"
            img.save(outpath)
            self.rotated = "not rotated"
        return PeaImage(outname, outpath, self.config_file)

    def affirm_image_is_portrait(self):
        """
        Function rotates the provided disk image if needed, otherwise the disk_image provided is returned.
        :return: rotated disk_image or just disk_image if already landscape.
        """
        img = Image.open(self.path)
        dimensions = img.size
        img.load()
        if dimensions[1] < dimensions[0]:
            img_landscape = img.transpose(PIL.Image.ROTATE_270)
            img_landscape.save(self.path)
        # add success for this process to the config


    @staticmethod
    def get_greyscale_copy(in_name, in_path, config_file):
        """
        Function: Gets a copy of this image transformed to grayscale.
        :param in_name: The name of the image to be copied and transformed.
        :return: A copy of the provided image in greyscale.
        """
        # Open image this image.
        input_image = cv2.cv.LoadImage(in_path)
        # Make a copy of same dimension.
        grey_image = cv2.cv.CreateImage(cv2.cv.GetSize(input_image), cv2.cv.IPL_DEPTH_8U, 1)
        # Transformation - Grey scale.
        cv2.cv.CvtColor(input_image, grey_image, cv2.cv.CV_BGR2GRAY)  # changed from cv.CV_BGR2GRAY
        # Make the save path with "grey.jpg" suffix.
        out_name, out_path = PeaUtils.get_file_path("B", in_path, ".jpg")
        # Save transformed image.
        cv2.cv.SaveImage(out_path, grey_image)
        # return a new instance of PeaImage that holds/represents the new greyscale copy of this image.
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    @staticmethod
    def get_adpative_threshold_copy(in_name, in_path, config_file):
        """
        Function: Gets a copy of this image transformed with adaptive threshold.
        :param retname:
        :param retpath:
        :return: An adaptive thresholded image
        """
        # Apply adaptive threshold
        # Open image this image.
        input_image = cv2.cv.LoadImage(in_path)
        # Make a copy of same dimension
        threshold_image = cv2.cv.CreateImage(cv2.cv.GetSize(input_image), cv2.cv.IPL_DEPTH_8U, 1)
        cv2.cv.AdaptiveThreshold(threshold_image, threshold_image, 255, cv2.cv.CV_ADAPTIVE_THRESH_GAUSSIAN_C,
                                 cv2.cv.CV_THRESH_BINARY, 101, 5)
        # Make the save path with "threshold.jpg" suffix.
        out_name, out_path = PeaUtils.get_file_path("C", in_path, ".jpg")
        # Save transformed image.
        cv2.cv.SaveImage(out_path, threshold_image)
        # return a new instance of PeaImage adaptive thresholded greyscale copy of this image.
        # add success for this process to the config
        return PeaImage(out_name, out_path,config_file)

    @staticmethod
    def get_cropped_copy(in_name, in_path, percent_x, percent_y, config_file):
        """
        Function : Crop the image to the specified size by parameters provided
        :param in_name: Name of image to be cropped
        :param in_path: Path of image to be cropped
        :param percent_x: Percent of the image's width within the crop
        :param percent_y: Percent of the image's height within the crop
        :return: Copy of the cropped region where the region is in proportion to the dimension.
        """
        input_image = Image.open(in_path)
        dimension = input_image.size
        right = (dimension[0] / 100) * percent_x
        lower = (dimension[1] / 100) * percent_y
        input_image = input_image.crop((0, 0, right, lower))
        temp = input_image.copy()
        out_name, out_path = PeaUtils.get_file_path("D", in_path, ".jpg")
        temp.save(out_path)
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    @staticmethod
    def get_denoised_copy(in_name, in_path, config_file):
        """
        Function: Denoises the label
        :param retname:
        :param retpath:
        :add param descriptions
        :return: Denoised label
        """
        input_image = cv2.imread(in_path)
        # Transformation - denoise the image
        denoised = cv2.fastNlMeansDenoising(input_image, None, 50, 21, 7)  # TODO: un hard code
        # Make save path with "denoised.jpg" suffix
        out_name, out_path = PeaUtils.get_file_path("E", in_path, ".jpg")
        cv2.imwrite(out_path, denoised)
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    @staticmethod
    def remove_intermediate_file(file_to_be_removed):
        """
        Function: Remove unwanted files
        :param file_to_be_removed:
        :return:
        """
        os.remove(file_to_be_removed)

    def delete_myself_from_disk(self):
        """
        Function: Removes the file for this image object from disk
        """
        os.remove(self.path)

    def resize_image(self, ret_name, ret_path):
        """
        Function: Resize images by 50% more pixels
        :param ret_path:
        :param ret_name:
        :param jpgpathtofile:
        :return: A copy of this resized image.
        """
        # Open the file
        reference_image = io.imread(self.path)
        # Transformation - resize by 50%
        pea_resize = rescale(reference_image, 0.5, mode='constant')  # 50% more pixels
        # Create Path
        save_path = self.path + ".R.jpg"
        # Save Image
        pea_resize.save(save_path, pea_resize)
        # return pathtoresizedimage
        # add success for this process to the config
        self.config_file.R_resize_output_config_file(save_path)
        return PeaImage(ret_name, ret_path, self.config_file)

    @staticmethod
    def resize_image_pc(img_path, pc):
        reference_image = io.imread(img_path)
        resized_img = rescale(reference_image, pc, mode='constant')
        io.imsave(img_path + ".resized_word.jpg", resized_img)
        return img_path + ".resized_word.jpg"

    def convert_to_hsv(self, ret_name, ret_path):
        """
        Function: Convert colour space to hsv
        :param ret_path:
        :param ret_name:
        :param pathtoresizedimage:
        :return:A copy of this image in hsv.
        """
        # Open the file
        reference_image = io.imread(self.path)
        # Make a copy
        hsv_image = reference_image
        # Transformation - convert to HSV
        pea_hsv = color.rgb2hsv(hsv_image)  # convert to hsv
        # Create Path
        save_path = self.path + ".H.jpg"
        # Save image
        pea_hsv.save(save_path, pea_hsv)
        # add success for this process to the config
        self.config_file.H_hsv_output_config_file(save_path)
        # return pathtohsvimage
        return PeaImage(ret_name, ret_path, self.config_file)

    def get_hsv_value(self):
        """
        Function: Extract Value from HSV file. That is all hue dimension, all sat dimensions, only 2nd index in brightness value
        :param pathtohsvimage:
        :return:
        """
        # Open the file
        reference_image = io.imread(self.path)
        # Make a copy
        extract_value = reference_image
        # Extract value
        # every thing in hue dimension, everything in sat, only 2nd index in brightness value
        self.hsv_value = extract_value[:, :, 2]
        return self.hsv_value

    def get_hue_value(self):
        # Open the file
        reference_image = io.imread(self.path)
        # Make a copy
        extract_hue_value = reference_image
        # Extract hue value
        # every thing in hue dimension, everything in sat, only 0th index in brightness value
        self.hue_value = self.hsv_value[:, :, 0]
        return self.hue_value

    def get_shape(self, pea_hsv):
        """
        Find the shape of the hsv image, x by y in pixels, and how many hsv values each pixel has
        :param pea_hsv:
        :return: self.shape
        """
        self.shape = pea_hsv.shape  # find geometry of shape, x by y in pixels, and how many hsv values each pixel has
        return self.shape

    def get_mean_hsv(self, pea_resize):
        """
        Function: find mean aggregate HSV of resized image
        :param pea_resize:
        :return:
        """
        self.mean_hsv = pea_resize.mean()
        return self.mean_hsv

    def get_mean_hue_value(self, pea_value):
        """
        Function: find mean aggregate hue valued extracted pea shapes
        :param pea_value:
        :return:
        """
        self.mean_hue_value = pea_value.mean()
        return self.mean_hue_value

    def get_median_resize(self, pea_resize):
        """
         Function: find median HSV of resized image
         :param pea_resize:
         :return:
         """
        self.median_resize = np.median(pea_resize)
        return self.median_resize

    def get_median_hue_value(self, pea_hue_value):
        """
        Function: find median HSV of valued extracted pea shapes
        :param pea_hue_value:
        :return:
        """
        self.median_hue_value = np.median(pea_hue_value)
        return self.median_hue_value

    def get_hue_value_thresholding(self, pea_hue_value, median_hue_value):
        """
        Function: Value thresholding using median intensity value what becomes white/black
        if pixel colour is above median value is one/true then white
        if pixel colour is below median value is zero/false then black
        Scale_ROI is true or false of Pea_Hue matrix
        0.925 Hard-coded, should use Python Stats to calculat
        :param pea_hue_value:
        :param median_hue_value:
        :return:
        """
        self.scale_ROI = pea_hue_value > median_hue_value * 0.925
        return self.scale_ROI

    def get_edge_detection(self, pea_hue_value):
        """
        Function: Canny edge detection to find the outline of the leaf
        input greyscale image, find the edge and do not smooth with sigma
        :param pea_hue_value:
        :return:self.edge_detection
        """
        self.edge_detection = feature.canny(color.rgb2gray(pea_hue_value), sigma=3)
        return self.edge_detection

    def get_dilation(self, edge_detection):
        """
        Function: Dilate edges
        dilation takes the pixels of the edge and sets its neighbour to be brighter, disk takes 1 pixel in each direction, changes thicknes
        :param edge_detection:
        :return:
        """
        self.dilated_Edge = dilation(edge_detection, disk(3))
        return self.dilated_Edge

    def inverse_roi(self, scale_roi):
        """
        Function: # create inverse of Scale ROI
        draw white if the two pictures above differ, black if they are the same
        :param scale_roi:
        :return:
        """
        self.refined_ROI = np.logical_not(scale_roi)
        return self.refined_ROI

    def find_objects(self, peahue_value):
        """
        Function: Find objects
        # Close objects to fill holes in detected objects
        # Find contours at a constant value of 0.9
        :param peahue_value:
        :return: newRegionsSorted
        """
        leaf_mask = self.refined_ROI
        # How many neighbours does it have with only connected by 2
        image_labels = measure.label(leaf_mask, connectivity=2)  # 2D matrix of connected pixels
        # begin taking measurements
        # co-ordinates of connected pixels and value
        regions = regionprops(image_labels, intensity_image=self.hsv_value)
        # Set empty array and variables for storing measurements SECOND PEA PROBLEM STARTS
        blank_img = np.zeros((peahue_value[0], peahue_value.shape[1]), dtype=np.uint8)
        largest_area = 0
        self.leaf_tmp = blank_img.copy()  # also another array of zeros
        self.contour_tmp = blank_img.copy()  # also another array of zeros
        # creates a new sorted region by area
        self.newRegionsSorted = sorted(regions, key=attrgetter('area'), reverse=True)
        return PeaImage(self.newRegionsSorted, self.leaf_tmp, self.contour_tmp, self.config_file)

    @staticmethod
    def global_thresholding(in_name, in_path, config_file):
        img = cv2.imread(in_path, 0)
        # global thresholding
        ret1, th1 = cv2.threshold(img, 127, 255, cv2.THRESH_BINARY)
        out_name, out_path = PeaUtils.get_file_path("F", in_path, ".jpg")
        cv2.imwrite(out_path, th1)
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    @staticmethod
    def otsus_thresholding(in_name, in_path, config_file):
        img = cv2.imread(in_path, 0)
        # Otsu's thresholding
        ret2, th2 = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        out_name, out_path = PeaUtils.get_file_path("G", in_path, ".jpg")
        cv2.imwrite(out_path, th2)
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    def get_loaded_image(self):
        img = cv2.imread(self.path, 0)
        return img

    @staticmethod
    def gaussian_filtering_otsus_thresholding(in_name, in_path, config_file):
        img = cv2.imread(in_path, 0)
        # Otsu's thresholding after Gaussian filtering
        blur = cv2.GaussianBlur(img, (5, 5), 0)
        ret3, th3 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        out_name, out_path = PeaUtils.get_file_path("H", in_path, ".jpg")
        cv2.imwrite(out_path, th3)
        # add success for this process to the config
        return PeaImage(out_name, out_path, config_file)

    @staticmethod
    def get_line_length(p1, p2):
        a = [p1[0] - p2[0], p1[1] - p2[1]]
        b = math.sqrt(a[0] ** 2 + a[1] ** 2)
        return b

    @staticmethod
    def get_unit_vector(p1, p2):
        """
        Get unit vector of two points
        :param p1:
        :param p2:
        :return:
        """
        distance = [p1[0] - p2[0], p1[1] - p2[1]]
        norm = math.sqrt(distance[0] ** 2 + distance[1] ** 2)
        direction = [distance[0] / norm, distance[1] / norm]
        return direction

        # ###############
        # # for teeth convexity defects
        # max_sinus_distance = w / 4.0
        # hull = cv2.convexHull(contour, returnPoints=False)
        # defects = cv2.convexityDefects(contour, hull)
        # gray_for_drawing_7 = cv2.cvtColor(gray, cv2.COLOR_GRAY2BGR)
        # teeth_count = 0
        # ######################
        # # unit vector boundary testing.
        # horizontal_direction = self.get_unit_vector(centroid, (centroid[0] + 40, centroid[1]))
        # print "horizontal direction: ", horizontal_direction
        # cv2.line(gray_for_drawing_7, centroid, (centroid[0] + 40, centroid[1]), [0, 255, 0], 1)
        # cv2.circle(gray_for_drawing_7, centroid, 1, [255, 0, 0], -1)
        # cv2.circle(gray_for_drawing_7, (centroid[0] + 40, centroid[1]), 1, [0, 0, 255], -1)
        # up_vertical = self.get_unit_vector(centroid, (centroid[0], centroid[1] - 40))
        # print "up vertical: ", up_vertical
        # down_vertical = self.get_unit_vector(centroid, (centroid[0], centroid[1] + 40))
        # print "down vertical: ", down_vertical
        # ######################
        # for i in range(defects.shape[0]):
        #     s, e, f, d = defects[i, 0]
        #     start = tuple(contour[s][0])
        #     end = tuple(contour[e][0])
        #     far = tuple(contour[f][0])
        #     # sinus_length = cv2.norm(np.asarray(start),np.asarray(end))
        #     # print "sinus length", sinus_length
        #     tooth_width = cv2.norm(np.asarray(start), np.asarray(far))
        #     if tooth_width > 2:
        #         teeth_count += 1
        #         if far[0] >= centroid[0]:
        #             print "start, far: ", start, far
        #             direction = self.get_unit_vector(start, far)
        #             print "unit vector: ", direction
        #             # cv2.line(gray_for_drawing_7, start, end, [0, 0, 0], 1)
        #             cv2.line(gray_for_drawing_7, start, far, [255, 255, 255], 1)
        #             cv2.line(gray_for_drawing_7, end, far, [0, 255, 0], 1)
        #             cv2.circle(gray_for_drawing_7, far, 1, [0, 0, 255], -1)
        #             cv2.circle(gray_for_drawing_7, start, 1, [255, 0, 0], -1)
        #             # cv2.circle(gray_for_drawing_7, end, 1, [255,255,255], -1)
        #         else:
        #             print "end, far: ", end, far
        #             direction = self.get_unit_vector(end, far)
        #             print "unit vector: ", direction
        #             # cv2.line(gray_for_drawing_7, start, end, [0, 0, 0], 1)
        #             cv2.line(gray_for_drawing_7, end, far, [255, 255, 255], 1)
        #             cv2.line(gray_for_drawing_7, start, far, [0, 255, 0], 1)
        #             cv2.circle(gray_for_drawing_7, far, 1, [0, 0, 255], -1)
        #             cv2.circle(gray_for_drawing_7, end, 1, [255, 0, 0], -1)
        #             # cv2.circle(gray_for_drawing_7, end, 1, [255,255,255], -1)
        # cv2.putText(gray_for_drawing_7, str(teeth_count), centroid, cv2.FONT_HERSHEY_SIMPLEX, 2, 255)
        # cv2.imwrite(path + ".teeth.jpg", gray_for_drawing_7)
        # #################
